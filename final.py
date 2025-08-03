from flask import Flask, render_template, request, send_file, redirect, url_for, flash, make_response, jsonify, session
import pandas as pd
import os
import subprocess
import time
import pyautogui
from functools import wraps
from werkzeug.utils import secure_filename
import win32com.client
from docx import Document
from datetime import datetime, timedelta
from selenium import webdriver
from selenium.webdriver.chrome.service import Service
from selenium.webdriver.common.by import By
from selenium.webdriver.chrome.options import Options
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
from selenium.webdriver.support.ui import Select
from webdriver_manager.chrome import ChromeDriverManager
import threading
import logging
import re
import json
import requests

app = Flask(__name__)
app.secret_key = 'your_secret_key_change_this_in_production'
app.config['UPLOAD_FOLDER'] = 'uploads/'
app.config['MAX_CONTENT_LENGTH'] = 25 * 1024 * 1024  # 25MB max file size

# Hardcoded download filename and path
app.config['DOWNLOAD_FILENAME'] = 'ROB.xlsx'
app.config['DOWNLOAD_PATH'] = r'C:\Users\vishwas\Desktop\RPA\\' + app.config['DOWNLOAD_FILENAME']

# Ensure upload folder exists
os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)

# Set up logging to capture output
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Add a global variable to track processing status
processing_status = {
    'active': False,
    'message': 'Ready',
    'progress': 0,
    'total': 0,
    'current_file': '',
    'logs': []
}

def allowed_file(filename):
    """Check if file extension is allowed"""
    ALLOWED_EXTENSIONS = {'xlsx', 'xls', 'csv'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def allowed_backend_file(filename):
    """Check if backend file extension is allowed"""
    BACKEND_EXTENSIONS = {'xlsx', 'xls'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in BACKEND_EXTENSIONS

def log_to_status(message):
    """Add a message to the processing status logs"""
    global processing_status
    processing_status['logs'].append(f"{datetime.now().strftime('%H:%M:%S')}: {message}")
    print(f"[LOG] {message}")

# ============================================================================
# HOME ROUTE
# ============================================================================

@app.route('/')
def home():
    return render_template('home.html')
@app.route('/index.html')
def index():
    return render_template('index.html')

# ============================================================================
# DOCUMENT PROCESSING ROUTES
# ============================================================================
@app.route('/document_processing', methods=['GET', 'POST'])
def document_processing():
    global processing_status
    
    if request.method == 'POST':
        try:
            # Get form data - use session data as defaults if available
            article_code = request.form.get('article_code') or request.form.get('open_pr_id') or session.get('open_pr_id', '6HA-2025-M6K439')
            author_name = request.form.get('author_name') or session.get('username', 'vishwas tiwari')
            author_email = request.form.get('author_email') or session.get('email', 'vishwas@coherentmarketinsights.com')
            company_name = request.form.get('company_name', 'Coherent Market Insights')
            phone_number = request.form.get('phone_number') or session.get('mobile', '1234567890')
            
            # Get all 4 image paths from form
            image_path1 = request.form.get('image_path1') 
            image_path2 = request.form.get('image_path2') 
            image_path3 = request.form.get('image_path3') 
            image_path4 = request.form.get('image_path4')
            image_path5 = request.form.get('image_path5')
            image_path6 = request.form.get('image_path6')
            image_path7 = request.form.get('image_path7')
            image_path8 = request.form.get('image_path8')
            image_path9 = request.form.get('image_path9')
            image_path10 = request.form.get('image_path10')
            
            # Create list of image paths (only include non-empty paths)
            image_paths = []
            for path in [image_path1, image_path2, image_path3, image_path4,image_path5,image_path6,image_path7,image_path7,image_path8,image_path9,image_path10]:
                if path and path.strip():
                    image_paths.append(path.strip())
            
            # Power Automate output folder path
            custom_folder = request.form.get('custom_folder')
            if custom_folder:
                folder_path = custom_folder
            else:
                today = datetime.today()
                folder_path = rf'C:\Users\vishwas\Desktop\RPA\Files\{today.year}\{today.strftime("%m")}\{today.strftime("%d")}'
            
            processing_mode = request.form.get('processing_mode', 'manual')
            
            # Validate paths before processing
            excel_path = r'C:\Users\vishwas\Desktop\RPA\ROB.xlsx'
            
            # Check if required files exist
            validation_errors = []
            if not os.path.exists(excel_path):
                validation_errors.append(f"Excel file not found: {excel_path}")
            if not os.path.exists(folder_path):
                validation_errors.append(f"Folder not found: {folder_path}")
            
            # Validate image paths
            if not image_paths:
                validation_errors.append("At least one image path is required")
            else:
                for i, path in enumerate(image_paths, 1):
                    if not os.path.exists(path):
                        validation_errors.append(f"Image file {i} not found: {path}")
            
            if validation_errors:
                for error in validation_errors:
                    flash(error)
                return render_template('document_processing.html', 
                                     session_data={
                                         'username': session.get('username', ''),
                                         'email': session.get('email', ''),
                                         'mobile': session.get('mobile', ''),
                                         'open_pr_id': session.get('open_pr_id', ''),
                                         'image_path1': image_path1 or '',
                                         'image_path2': image_path2 or '',
                                         'image_path3': image_path3 or '',
                                         'image_path4': image_path4 or '',
                                         'image_path5': image_path5 or '',
                                         'image_path6': image_path6 or '',
                                        'image_path7': image_path7 or '',
                                        'image_path8': image_path8 or '',
                                        'image_path9': image_path9 or '',
                                         'image_path10': image_path10 or ''
                                     })
            
            # Reset processing status
            processing_status = {
                'active': True,
                'message': 'Starting processing...',
                'progress': 0,
                'total': 0,
                'current_file': '',
                'logs': []
            }
            
            # Start processing in background thread - NOW INCLUDING image_paths list
            if processing_mode == 'auto':
                threading.Thread(target=process_documents_auto_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_paths)).start()
            else:
                threading.Thread(target=process_documents_manual_with_feedback, 
                               args=(folder_path, article_code, author_name, author_email, 
                                    company_name, phone_number, image_paths)).start()
            
            flash('Processing started! Check the status page for updates.')
            return redirect(url_for('processing_status'))
            
        except Exception as e:
            flash(f'Error starting processing: {str(e)}')
            logger.error(f"Error in document_processing: {e}")
            return render_template('document_processing.html', 
                                 session_data={
                                     'username': session.get('username', ''),
                                     'email': session.get('email', ''),
                                     'mobile': session.get('mobile', ''),
                                     'open_pr_id': session.get('open_pr_id', ''),
                                     'image_path1': request.form.get('image_path1', ''),
                                     'image_path2': request.form.get('image_path2', ''),
                                     'image_path3': request.form.get('image_path3', ''),
                                     'image_path4': request.form.get('image_path4', ''),
                                     'image_path5': request.form.get('image_path5', ''),
                                    'image_path6': request.form.get('image_path6', ''),
                                    'image_path7': request.form.get('image_path7', ''),
                                    'image_path8': request.form.get('image_path8', ''),
                                    'image_path9': request.form.get('image_path9', ''),
                                    'image_path10': request.form.get('image_path10', '')

                                 })
    
    # Pre-populate form with session data if available
    return render_template('document_processing.html', 
                         session_data={
                             'username': session.get('username', ''),
                             'email': session.get('email', ''),
                             'mobile': session.get('mobile', ''),
                             'open_pr_id': session.get('open_pr_id', ''),
                             'image_path1': session.get('image_path1', ''),
                             'image_path2': session.get('image_path2', ''),
                             'image_path3': session.get('image_path3', ''),
                             'image_path4': session.get('image_path4', '')
                         })

@app.route('/processing_status')
def processing_status_page():
    return render_template('processing_status.html')

@app.route('/api/get_processing_status')
def get_processing_status():
    """API endpoint to get current processing status"""
    global processing_status
    return jsonify(processing_status)

# ============================================================================
# DOCUMENT PROCESSING FUNCTIONS
# ============================================================================
import win32com.client
import re

import win32com.client
import re

def text_of_press_release(doc_path):
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in the background

    # Open the .doc file (adjust the file path if necessary)
    doc2 = word.Documents.Open(doc_path)

    # Extract the entire text from the document
    doc_text = doc2.Content.Text

    # Remove the first line from the document text
    lines = doc_text.splitlines()
    if len(lines) > 1:
        doc_text = '\n'.join(lines[1:])

    # Define the headings for which you want to add line breaks
    headings = [
        "➤Market Size and Overview",
        "➤Actionable Insights",
        "➤Actionable insights",
        "➤Growth factors",
        "➤Growth Factors",
        "➤Market trends",
        "➤Market Trends",
        "➤Key takeaways ",
        "➤Key Takeaways",
        "➤Market Segment and Regional Coverage ",
        "➤Market segment and regional coverage",
        "➤Key players",
        "➤Key Players",
        "➤Competitive Strategies and Outcomes",
        "❓ Frequently Asked Questions",
        "❓ Frequently asked questions"
    ]

    # FIXED: Add a line space AFTER each heading (not before and after)
    for heading in headings:
        doc_text = doc_text.replace(heading, f"{heading}\n")

    # Define the regex pattern for URLs
    url_pattern = re.compile(r"(https?://[^\s]+)")
    
    # Define regex patterns for FAQ questions (numbered questions and roman numerals)
    faq_pattern_numbers = re.compile(r"^\d+\.\s")  # Matches "1. ", "2. ", etc.
    faq_pattern_roman = re.compile(r"^[ivxlcdmIVXLCDM]+\.\s")  # Matches "i. ", "ii. ", "I. ", "II. ", etc.
    
    # Define regex pattern for CTA links (➔)
    cta_pattern = re.compile(r"^➔")  # Matches lines starting with ➔

    # Split the text into lines
    lines = doc_text.splitlines()
    processed_lines = []

    # Iterate over each line
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines in processing, we'll add them strategically
        if not line_stripped:
            continue
            
        # Check if this is a CTA line
        is_cta = cta_pattern.match(line_stripped)
        
        # Check if previous line was a CTA (for adjacent CTA handling)
        prev_was_cta = False
        if processed_lines:
            last_non_empty = None
            for prev_line in reversed(processed_lines):
                if prev_line.strip():
                    last_non_empty = prev_line.strip()
                    break
            if last_non_empty and cta_pattern.match(last_non_empty):
                prev_was_cta = True
        
        # Check if this line is a heading (starts with ➤ or ❓)
        is_heading = line_stripped.startswith('➤') or line_stripped.startswith('❓')
        
        # If a line contains a URL, add space before and after the URL
        if url_pattern.search(line):
            # Add space before (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            processed_lines.append('')  # Add space after
            
        # If a line is an FAQ question (starts with number or roman numeral), add space before it
        elif faq_pattern_numbers.match(line_stripped) or faq_pattern_roman.match(line_stripped):
            # Add space before FAQ question (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this is a CTA line
        elif is_cta:
            # Add space before CTA (unless previous was also CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this line is a heading (starts with ➤ or ❓)
        elif is_heading:
            # Add space before heading (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            # FIXED: Add space AFTER heading
            processed_lines.append('')
            
        else:
            # Regular content line
            processed_lines.append(line)

    # Join the processed lines back into a single string
    chunk = "\n".join(processed_lines)
    
    # Clean up multiple consecutive empty lines (replace with single empty line)
    chunk = re.sub(r'\n\s*\n\s*\n+', '\n\n', chunk)

    # Close the document
    doc2.Close()
    word.Quit()

    # Return the processed content
    return chunk
# Add this function at the top of your file if you haven't already
def safe_str_convert(value):
    """Safely convert any value to string, handling NaN, None, and numpy types"""
    if pd.isna(value) or value is None:
        return ''
    return str(value).strip()

# Replace the problematic section in run_selenium_automation_single function
# Find this section and replace it:

def run_selenium_automation_single(row_data, category, article_code, author_name, author_email, company_name, phone_number, image_paths, market_name):
    """Run Selenium automation for a single press release submission"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
    """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.
 """    
    """ Author of this marketing PR :
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. 
"""    
    """ Author of this marketing PR:

Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.
"""]
        
        # RANDOMLY SELECT IMAGE PATH from the provided list
        if image_paths and len(image_paths) > 0:
            selected_image_path = random.choice(image_paths)
            log_to_status(f"Randomly selected image: {selected_image_path}")
        else:
            log_to_status("No image paths provided!")
            return False, "No image paths available"
        
        # FIXED: Extract data from the row using safe_str_convert
        market_name = safe_str_convert(row_data.get('Market Name', ''))
        category = safe_str_convert(row_data.get('Category', ''))
        
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        # FIXED: Extract companies covered using safe_str_convert
        companies_raw = row_data.get('Companies covered', '')
        companies = safe_str_convert(companies_raw)
        log_to_status(f"Trying to process companies: '{companies}' for market '{market_name}'")

        # Create article title from market name and companies - FIXED
        if companies and companies.strip():  # companies is already a string now
            try:
                company_list = [c.strip() for c in companies.split(',') if c.strip()]
                first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
            except Exception as e:
                log_to_status(f"Error processing companies: {e}")
                first_five_companies = 'No companies available'
        else:
            first_five_companies = 'No companies available'
            
        random_prompt = random.choice(TITLE_PROMPTS)
        article_title = f"{market_name} {random_prompt} {first_five_companies}"
        
        # FIXED: Create multiline text from the row data using safe_str_convert
        market_size = safe_str_convert(row_data.get('Market Size', 'Market analysis and insights'))
        forecast_period = safe_str_convert(row_data.get('Forecast Period', 'N/A'))  
        cagr = safe_str_convert(row_data.get('CAGR', 'N/A'))
        key_players = safe_str_convert(row_data.get('Key Players', 'Leading companies in the market'))
        
        multiline_text = f"""
{market_name} - Market Insights Report

Market Overview:
{market_size}

Forecast Period: {forecast_period}
CAGR: {cagr}

Key Market Players:
{key_players}

For more detailed information, please refer to our comprehensive market research report.
        """
        
        log_to_status(f"Processing: {market_name}")
        log_to_status(f"Using category: {category}")

        log_to_status("Starting Selenium automation for: " + market_name)
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # Handle category selection with better error handling
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        
        # Wait a moment for the dropdown to be fully loaded
        time.sleep(1)
        
        # Get all available options from the dropdown
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        
        # FIXED: Use the category directly from the row data with fallback
        website_category = category if category else "Business,Economy,Finance,Banking & Insurance"
        log_to_status(f"Trying to select category: '{website_category}'")
        
        # Check if the exact category exists in the dropdown
        if website_category in available_options:
            log_to_status(f"Category '{website_category}' found in dropdown options")
        else:
            log_to_status(f"Category '{website_category}' NOT found in dropdown options")
            log_to_status(f"Available options are: {available_options}")
        
        try:
            # Try different selection methods
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}' by visible text: {e}")
            
            # Try selecting by value if visible text fails
            try:
                for option in select_obj.options:
                    if option.text.strip() == website_category:
                        select_obj.select_by_value(option.get_attribute('value'))
                        log_to_status(f"Successfully selected category by value: '{website_category}'")
                        break
                else:
                    raise Exception(f"Could not find option with text '{website_category}'")
            except Exception as e2:
                log_to_status(f"Error selecting category by value: {e2}")
                
                # Final fallback - try to select "Business,Economy,Finance,Banking & Insurance" directly if we have a health-related category
                try:
                    if "health" in website_category.lower() or "medicine" in website_category.lower():
                        select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                        log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance' as fallback for health-related category")
                    else:
                        select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                        log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance' as fallback")
                except Exception as e3:
                    log_to_status(f"Final fallback also failed: {e3}")
                    select_obj.select_by_index(1)  # Select first real option
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        
        title.clear()
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

            Mr. Shah
            Coherent Market Insights
            533 Airport Boulevard,
            Suite 400, Burlingame,
            CA 94010, United States
            US: + 12524771362
            UK: +442039578553
            AUS: +61-8-7924-7805
            India: +91-848-285-0837"""
        about.send_keys(multi)
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        # USE THE RANDOMLY SELECTED IMAGE PATH
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(selected_image_path)  # ✅ Using randomly selected image
        log_to_status(f"Uploaded image: {selected_image_path}")
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        caption.send_keys("Market research report visualization and analysis.")
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("Comprehensive market insights and analysis report submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(5)
        log_to_status(f"Selenium automation completed successfully for: {market_name}")
        log_to_status(f"Used image: {selected_image_path}")
        driver.quit()  # Quit the browser tab after publication
        return True, market_name
        
    except Exception as e:
        market_name = safe_str_convert(row_data.get('Market Name', 'Unknown')) if 'row_data' in locals() else 'Unknown'
        log_to_status(f"Selenium automation error for {market_name}: {e}")
        try:
            driver.quit()
        except:
            pass
        return False, market_name


def run_selenium_automation_all_rows(article_code, author_name, author_email, company_name, phone_number, image_paths):
    """Run Selenium automation for all rows in the Excel file"""
    try:
        # Read all data from Excel
        import pandas as pd
        excel_path = r'C:\Users\vishwas\Desktop\RPA\ROB.xlsx'
        df = pd.read_excel(excel_path)
        
        log_to_status(f"Found {len(df)} rows in Excel file")
        log_to_status(f"Available images: {len(image_paths)} paths")
        
        # Validate image paths
        if not image_paths or len(image_paths) == 0:
            log_to_status("❌ No image paths provided!")
            return 0, len(df)
        
        # Log available images
        for i, path in enumerate(image_paths, 1):
            log_to_status(f"Image {i}: {path}")
        
        # Results tracking
        successful_submissions = []
        failed_submissions = []
        
        # Process each row
        for index, row in df.iterrows():
            log_to_status(f"\n{'='*50}")
            log_to_status(f"Processing Row {index + 1} of {len(df)}")
            log_to_status(f"{'='*50}")
            
            try:
                # FIXED: Use safe_str_convert for category and market_name
                category = safe_str_convert(row.get('Category', ''))
                market_name = safe_str_convert(row.get('Market Name', f'Row {index + 1}'))
                
                success, result_market_name = run_selenium_automation_single(
                    row_data=row,
                    category=category,
                    article_code=article_code,
                    author_name=author_name,
                    author_email=author_email,
                    company_name=company_name,
                    phone_number=phone_number,
                    image_paths=image_paths,
                    market_name=market_name
                )
                
                if success:
                    successful_submissions.append(result_market_name)
                    log_to_status(f"✅ Successfully submitted: {result_market_name}")
                else:
                    failed_submissions.append(result_market_name)
                    log_to_status(f"❌ Failed to submit: {result_market_name}")
                
                # Add a delay between submissions to avoid overwhelming the server
                '''if index < len(df) - 1:  # Don't wait after the last submission
                    log_to_status("Waiting 20 seconds before next submission...")
                    time.sleep(300) ''' # ✅ FIXED: Changed from 360 to 20 seconds as per your original code
                    
            except Exception as e:
                # FIXED: Use safe_str_convert for error handling
                market_name = safe_str_convert(row.get('Market Name', f'Row {index + 1}'))
                failed_submissions.append(market_name)
                log_to_status(f"❌ Error processing row {index + 1} ({market_name}): {e}")
                continue
        
        # Final summary
        log_to_status(f"\n{'='*50}")
        log_to_status("FINAL SUMMARY")
        log_to_status(f"{'='*50}")
        log_to_status(f"Total rows processed: {len(df)}")
        log_to_status(f"Successful submissions: {len(successful_submissions)}")
        log_to_status(f"Failed submissions: {len(failed_submissions)}")
        log_to_status(f"Images used: {len(image_paths)} different images randomly selected")
        
        if successful_submissions:
            log_to_status(f"\n✅ Successfully submitted:")
            for market in successful_submissions:
                log_to_status(f"  - {market}")
        
        if failed_submissions:
            log_to_status(f"\n❌ Failed submissions:")
            for market in failed_submissions:
                log_to_status(f"  - {market}")
        
        return len(successful_submissions), len(failed_submissions)
        
    except Exception as e:
        log_to_status(f"Error in run_selenium_automation_all_rows: {e}")
        return 0, 0

def run_selenium_automation(article_code, article_title, multiline_text, category, author_name, 
                          author_email, company_name, phone_number, image_paths, market_name):
    """Enhanced run_selenium_automation function that randomly selects from image_paths - FIXED VERSION"""
    try:
        import random
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:\nRavina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.\n """,
            """ Author of this marketing PR :\nMoney Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc. \n""",
            """ Author of this marketing PR:\n\nAlice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from Openpr her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights.\n"""
        ]
        
        # RANDOMLY SELECT IMAGE PATH from the provided list
        if image_paths and len(image_paths) > 0:
            selected_image_path = random.choice(image_paths)
            log_to_status(f"Randomly selected image: {selected_image_path}")
        else:
            log_to_status("No image paths provided!")
            return False
        
        log_to_status("Starting Selenium automation...")
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.openpr.com/')
        
        # Handle cookie consent
        try:
            reject = WebDriverWait(driver, 10).until(
                EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
            )
            reject.click()
        except:
            pass
        
        # Navigate to submit page
        submit = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
        )
        submit.click()
        
        # Enter article code
        input_box = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
        )
        input_box.clear()
        input_box.send_keys(article_code)
        
        # Submit code
        try:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
            )
            submit2.click()
        except:
            submit2 = WebDriverWait(driver, 5).until(
                EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
            )
            submit2.click()
        
        # Fill form fields
        name = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
        )
        name.send_keys(author_name)
        
        email = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
        )
        email.clear()
        email.send_keys(author_email)
        
        number = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
        )
        number.clear()
        number.send_keys(phone_number)
        
        ComName = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
        )
        ComName.clear()
        ComName.send_keys("Coherent Market Insights")
        
        s1 = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
        )
        s1.click()
        
        # FIXED: Handle category selection with safe string conversion
        Category_element = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
        )
        time.sleep(1)
        select_obj = Select(Category_element)
        available_options = [option.text.strip() for option in select_obj.options]
        log_to_status(f"Available dropdown options: {available_options}")
        
        # FIXED: Use safe_str_convert for category
        website_category = safe_str_convert(category) if category else "Business,Economy,Finance,Banking & Insurance"
        log_to_status(f"Trying to select category: '{website_category}'")
        
        try:
            select_obj.select_by_visible_text(website_category)
            log_to_status(f"Successfully selected category: '{website_category}'")
        except Exception as e:
            log_to_status(f"Error selecting category '{website_category}': {e}")
            try:
                select_obj.select_by_visible_text("Business,Economy,Finance,Banking & Insurance")
                log_to_status("Selected 'Business,Economy,Finance,Banking & Insurance'")
            except:
                select_obj.select_by_index(1)
                log_to_status("Selected first available option as final fallback")
        
        title = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
        )
        title.clear()
        log_to_status(f"Setting title: '{article_title}'")
        title.send_keys(article_title)
        
        text = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
        )
        text.clear()
        text.send_keys(multiline_text)
        
        about = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
        )
        about.clear()
        multi = """Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837
"""
        about.send_keys(multi)
        
        address = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
        )
        address.clear()
        random_author = random.choice(AUTHOR_DESCRIPTIONS)
        address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."

        address.send_keys(address_content)
        
        # USE THE RANDOMLY SELECTED IMAGE PATH
        image = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
        )
        image.clear()
        image.send_keys(selected_image_path)  # ✅ Using randomly selected image
        log_to_status(f"Uploaded image: {selected_image_path}")
        
        caption = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
        )
        caption.clear()
        # FIXED: Use safe_str_convert for market_name in case it's also a numpy type
        safe_market_name = safe_str_convert(market_name) if market_name else "Market Analysis"
        caption.send_keys(safe_market_name)
        
        notes = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
        )
        notes.clear()
        notes.send_keys("Report submission.")
        
        # Agree to terms
        tick1 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
        )
        tick1.click()
        
        tick2 = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
        )
        tick2.click()
        
        # Submit form
        final = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
        )
        final.click()
        
        time.sleep(10)
        log_to_status("Selenium automation completed successfully")
        log_to_status(f"Used image: {selected_image_path}")
        # FIXED: Close driver properly
        '''try:
            driver.quit()
        except:
            pass'''
        return True
        
    except Exception as e:
        log_to_status(f"Selenium automation error: {e}")
        # FIXED: Close driver on error
        try:
            driver.quit()
        except:
            pass
        return False

def process_documents_auto_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number, image_paths):
    """Process documents automatically with status feedback"""
    global processing_status
    
    try:
        log_to_status(f"Starting auto processing. Folder: {folder_path}")
        log_to_status(f"Available images: {len(image_paths)} paths")
        
        excel_path = r'C:\Users\vishwas\Desktop\RPA\ROB.xlsx'
        import random
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0
        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        def refine_title_with_openai(title):
                    try:
                        import openai
                        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])
                        prompt = f"Please improve the grammar, structure, and readability of this press release title to make it more interesting and engaging for readers. Keep all original words intact - only rearrange, or adjust formatting as needed and there should be no parenthesis at end or start of title : '{title}'" 
                        response = client.chat.completions.create(
                            model=OPENAI_CONFIG['MODEL'],
                            messages=[
                                {"role": "system", "content": "You are an expert editor for press releases."},
                                {"role": "user", "content": prompt}
                            ],
                            max_tokens=60,
                            temperature=2.0
                        )
                        return response.choices[0].message.content.strip()
                    except Exception as e:
                        log_to_status(f"OpenAI title refinement error: {e}")
                        return title

        
        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Auto-processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                    
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:3]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                    
                random_prompt = random.choice(TITLE_PROMPTS)
                x = f"{market_name} {random_prompt}|{first_five_companies}"

                article_title = refine_title_with_openai(x)  # Your existing title refinement
                
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                
                # Run automation with image_paths list
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_paths,market_name)
                
                if success:
                    log_to_status(f"SUCCESS: Published {market_name}")
                    processed_count += 1
                else:
                    log_to_status(f"FAILED: Could not publish {market_name}")
                
                time.sleep(20)  # Wait 10 seconds between submissions
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Auto-processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Auto processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Auto processing error: {e}")



def process_documents_manual_with_feedback(folder_path, article_code, author_name, author_email, company_name, phone_number,image_path):
    """Process documents with manual intervention and status feedback"""
    global processing_status
    import random
    
    try:
        log_to_status(f"Starting manual processing. Folder: {folder_path}")
        
        excel_path = r'C:\Users\vishwas\Desktop\RPA\ROB.xlsx'
        
        # Load Excel file
        log_to_status("Loading Excel file...")
        keywords_df = pd.read_excel(excel_path)
        market_names = keywords_df['Market Name'].dropna().tolist()
        
        processing_status['total'] = len(market_names)
        log_to_status(f"Found {len(market_names)} market names to process")
        
        processed_count = 0

        TITLE_PROMPTS = [
    "Is Booming Worldwide 2025-2032",
    "Generated Opportunities, Future Scope 2025-2032",
    "Future Business Opportunities 2025-2032",
    "Growth in Future Scope 2025-2032",
    "Is Booming So Rapidly Growth by 2032",
    "Is Booming So Rapidly with CAGR of 6.9%",
    "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
    "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
    "Set to Witness Significant Growth by 2025-2032",
    "to Witness Massive Growth by 2032",
    "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
    "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
    "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
    "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
    "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
    "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
]

        
        for i, market_name in enumerate(market_names):
            processing_status['progress'] = i
            processing_status['current_file'] = market_name
            processing_status['message'] = f"Processing {i+1} of {len(market_names)}: {market_name}"
            
            doc_file = f"ROB_{market_name}.doc"
            doc_path = os.path.join(folder_path, doc_file)
            
            log_to_status(f"Looking for file: {doc_path}")
            
            if os.path.exists(doc_path):
                log_to_status(f"Processing: {market_name}")
                
                # Extract text
                processing_status['message'] = f"Extracting text from {market_name}..."
                multiline_text = text_of_press_release(doc_path)
                # Get companies for this market from Excel row
                matching_row = keywords_df[keywords_df['Market Name'] == market_name]
                if not matching_row.empty:
                    companies = matching_row.iloc[0].get('Companies covered', '')
                else:
                    companies = ''
                if companies and isinstance(companies, str) and companies.strip():
                    company_list = [c.strip() for c in companies.split(',') if c.strip()]
                    first_five_companies = ', '.join(company_list[:5]) if company_list else 'No companies available'
                else:
                    first_five_companies = 'No companies available'
                random_prompt = random.choice(TITLE_PROMPTS)
                article_title = f"{market_name} {random_prompt} {first_five_companies}"
                
                # Get category for this market from Excel row
                category = matching_row.iloc[0].get('Category', '') if not matching_row.empty else ''
                # Run automation
                processing_status['message'] = f"Submitting {market_name} via automation..."
                success = run_selenium_automation(article_code, article_title, multiline_text, category, 
                                                author_name, author_email, company_name, phone_number, image_path)
                
                if success:
                    log_to_status(f"Published {market_name}")
                    processed_count += 1
                
                time.sleep(5)
                
            else:
                log_to_status(f"ERROR: File not found: {doc_path}")
        
        processing_status['active'] = False
        processing_status['message'] = f"Processing complete! Published {processed_count} of {len(market_names)} articles"
        log_to_status(f"Manual processing complete. Published {processed_count} articles.")
        
    except Exception as e:
        processing_status['active'] = False
        processing_status['message'] = f"Error: {str(e)}"
        log_to_status(f"EXCEPTION: Manual processing error: {e}")

# ============================================================================
# ROB PROCESSING ROUTES
# ============================================================================
@app.route('/rob', methods=['GET', 'POST'])
def rob():
    if request.method == 'POST':
        # Get user information
        username = request.form.get('username')
        email = request.form.get('email')
        open_pr_id = request.form.get('open_pr_id')
        mobile = request.form.get('mobile')

        # Validate required user fields
        if not all([username, email, open_pr_id, mobile]):
            flash('All user information fields are required!')
            return redirect(request.url)

        # Get domain totals
        hc_total = int(request.form.get('hc_total', 0))
        cmfe_total = int(request.form.get('cmfe_total', 0))
        ict_total = int(request.form.get('ict_total', 0))
        
        total_words = hc_total + cmfe_total + ict_total
        
        if total_words == 0:
            flash('Please enter at least one domain with words greater than 0!')
            return redirect(request.url)

        # Get HC subdomain percentages
        hc_subdomains = {
            'medical_devices': int(request.form.get('hc_medical_devices_percent', 0)),
            'pharmaceutical': int(request.form.get('hc_pharmaceutical_percent', 0)),
            'biotechnology': int(request.form.get('hc_biotechnology_percent', 0)),
            'healthcare_it': int(request.form.get('hc_healthcare_it_percent', 0)),
            'clinical_diagnostic': int(request.form.get('hc_clinical_diagnostic_percent', 0)),
            'medical_imaging': int(request.form.get('hc_medical_imaging_percent', 0))
        }
        
        # Get CMFE subdomain percentages  
        cmfe_subdomains = {
            'advanced_materials': int(request.form.get('cmfe_advanced_materials_percent', 0)),
            'consumer_goods': int(request.form.get('cmfe_consumer_goods_percent', 0)),
            'food_beverages': int(request.form.get('cmfe_food_beverages_percent', 0)),
            'bulk_chemicals': int(request.form.get('cmfe_bulk_chemicals_percent', 0)),
            'specialty_chemicals': int(request.form.get('cmfe_specialty_chemicals_percent', 0)),
            'energy': int(request.form.get('cmfe_energy_percent', 0)),
            'packaging': int(request.form.get('cmfe_packaging_percent', 0)),
            'food_ingredients': int(request.form.get('cmfe_food_ingredients_percent', 0)),
            'polymers_resins': int(request.form.get('cmfe_polymers_resins_percent', 0)),
            'agrochemicals': int(request.form.get('cmfe_agrochemicals_percent', 0)),
            'cosmetic_ingredients': int(request.form.get('cmfe_cosmetic_ingredients_percent', 0)),
            'green_chemicals': int(request.form.get('cmfe_green_chemicals_percent', 0))
        }
        
        # Get ICT subdomain percentages
        ict_subdomains = {
            'info_comm_tech': int(request.form.get('ict_info_comm_tech_percent', 0)),
            'automotive_transport': int(request.form.get('ict_automotive_transport_percent', 0)),
            'industrial_automation': int(request.form.get('ict_industrial_automation_percent', 0)),
            'smart_technologies': int(request.form.get('ict_smart_technologies_percent', 0)),
            'semiconductors': int(request.form.get('ict_semiconductors_percent', 0)),
            'consumer_electronics': int(request.form.get('ict_consumer_electronics_percent', 0)),
            'aerospace_defense': int(request.form.get('ict_aerospace_defense_percent', 0)),
            'construction_engineering': int(request.form.get('ict_construction_engineering_percent', 0))
        }

        # Validate file upload
        file = request.files.get('file')
        if not file or file.filename == '':
            flash('Weekly RID file is required!')
            return redirect(request.url)

        if not allowed_file(file.filename):
            flash('Only Excel files (.xlsx, .xls) and CSV files are allowed!')
            return redirect(request.url)

        # Save uploaded file
        filename = secure_filename(file.filename)
        input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        file.save(input_path)

        # Store user data in session
        session['username'] = username
        session['email'] = email
        session['open_pr_id'] = open_pr_id
        session['mobile'] = mobile
        
        # Create domain configuration
        domain_config = {
            'HC': {
                'total': hc_total,
                'subdomains': hc_subdomains
            },
            'CMFE': {
                'total': cmfe_total,
                'subdomains': cmfe_subdomains
            },
            'ICT': {
                'total': ict_total,
                'subdomains': ict_subdomains
            }
        }
        
        return redirect(url_for('process_rob_with_domains', 
                                file_path=input_path,
                                username=username, 
                                email=email,
                                open_pr_id=open_pr_id, 
                                mobile=mobile,
                                domain_config=json.dumps(domain_config)))
    
    return render_template('rob.html')
@app.route('/process_rob_with_domains')
def process_rob_with_domains():
    """Fixed version with proper error handling and guaranteed return"""
    try:
        # Get parameters from request
        file_path = request.args.get('file_path')
        username = request.args.get('username')
        email = request.args.get('email')
        open_pr_id = request.args.get('open_pr_id')
        mobile = request.args.get('mobile')
        domain_config_json = request.args.get('domain_config')
        
        print(f"DEBUG: Received parameters:")
        print(f"  file_path: {file_path}")
        print(f"  username: {username}")
        print(f"  domain_config_json: {domain_config_json}")

        # Validate required parameters
        if not file_path:
            flash('❌ Missing file path')
            return redirect(url_for('rob'))
            
        if not os.path.exists(file_path):
            flash('❌ File not found')
            return redirect(url_for('rob'))
            
        if not domain_config_json:
            flash('❌ Missing domain configuration')
            return redirect(url_for('rob'))
            
        if not username:
            flash('❌ Missing username')
            return redirect(url_for('rob'))

        # Parse domain configuration safely
        try:
            domain_config = json.loads(domain_config_json)
            print(f"DEBUG: Parsed domain_config: {domain_config}")
        except (json.JSONDecodeError, TypeError) as e:
            print(f"ERROR: Failed to parse domain config: {e}")
            flash('❌ Invalid domain configuration')
            return redirect(url_for('rob'))
        
        # Process the file with domain/subdomain logic
        print("DEBUG: Starting file processing...")
        result = process_file_with_domain_subdomain_logic(file_path, domain_config)
        print(f"DEBUG: Processing result: {result}")
        
        if not result:
            flash('❌ Processing function returned no result')
            return redirect(url_for('rob'))
            
        if not result.get('success', False):
            error_msg = result.get('error', 'Unknown processing error')
            flash(f'❌ Error processing file: {error_msg}')
            return redirect(url_for('rob'))

        # Store results in session
        session['rob_file_path'] = result.get('rob_output_path')
        session['remaining_file_path'] = result.get('remaining_output_path')
        session['remaining_filename'] = result.get('remaining_filename')
        session['extraction_summary'] = result.get('extraction_summary', {})
        session['shortage_info'] = result.get('shortage_info', {})
        
        # Session variables for existing template
        total_extracted = result.get('total_extracted', 0)
        total_remaining = result.get('total_remaining', 0)
        
        session['extracted_count'] = total_extracted
        session['remaining_count'] = total_remaining
        session['total_count'] = total_extracted + total_remaining
        session['records_processed'] = total_extracted + total_remaining
        
        # Create success messages
        flash(f'✅ Successfully processed {total_extracted} words across all domains!')
        
        # Show domain distribution results
        extraction_summary = result.get('extraction_summary', {})
        for domain, stats in extraction_summary.items():
            extracted = stats.get('extracted', 0)
            if extracted > 0:
                flash(f'📊 {domain}: {extracted} words extracted')
        
        # Show shortage notifications
        shortage_info = result.get('shortage_info', {})
        if shortage_info.get('has_shortages', False):
            flash('⚠️ Some adjustments were made due to domain/subdomain shortages')
        
        # Get weekly_rid_folder for template
        weekly_rid_folder = result.get('weekly_rid_folder', 'Weekly_RID folder')
        
        print(f"DEBUG: Rendering template with:")
        print(f"  extracted_count: {total_extracted}")
        print(f"  remaining_count: {total_remaining}")
        print(f"  total_count: {total_extracted + total_remaining}")
        print(f"  username: {username}")
        
        # Render results page with your existing template
        return render_template('rob_result.html',
                             extracted_count=total_extracted,
                             remaining_count=total_remaining,
                             total_count=total_extracted + total_remaining,
                             username=username,
                             records_processed=total_extracted + total_remaining,
                             weekly_rid_folder=weekly_rid_folder,
                             remaining_filename=result.get('remaining_filename', 'remaining_file.xlsx'))

    except Exception as e:
        print(f"ERROR in process_rob_with_domains: {e}")
        import traceback
        traceback.print_exc()
        flash(f'❌ Unexpected error: {str(e)}')
        return redirect(url_for('rob'))
 


#Also update the main process function to handle the weekly_rid_folder
def process_file_with_domain_subdomain_logic(file_path, domain_config):
    """Process uploaded file with domain and subdomain extraction logic - FIXED VERSION"""
    try:
        print(f"DEBUG: Starting file processing with config: {domain_config}")
        
        # Read the uploaded file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        print(f"DEBUG: Original file has {total_rows} rows")
        print(f"DEBUG: Columns: {list(df_original.columns)}")
        
        # Validate required columns exist
        required_columns = ['Domain', 'Sub Domain']
        missing_columns = [col for col in required_columns if col not in df_original.columns]
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}. Expected columns: Domain, Sub D~omain'
            }
        
        # Calculate subdomain allocations for each domain
        allocation_plan = calculate_subdomain_allocations(domain_config)
        print(f"DEBUG: Allocation Plan: {allocation_plan}")
        
        # Extract rows according to allocation plan
        extraction_result = extract_rows_by_domain_subdomain(df_original, allocation_plan)
        
        if not extraction_result.get('success', False):
            return extraction_result
        
        extracted_df = extraction_result['extracted_df']
        remaining_df = extraction_result['remaining_df']
        extraction_summary = extraction_result['extraction_summary']
        shortage_info = extraction_result['shortage_info']
        
        print(f"DEBUG: Extraction completed. Extracted: {len(extracted_df)}, Remaining: {len(remaining_df)}")
        
        # Save results
        save_result = save_domain_extraction_results(extracted_df, remaining_df)
        
        if not save_result.get('success', False):
            return save_result
        
        return {
            'success': True,
            'total_extracted': len(extracted_df),
            'total_remaining': len(remaining_df),
            'extraction_summary': extraction_summary,
            'shortage_info': shortage_info,
            'rob_output_path': save_result['rob_output_path'],
            'remaining_output_path': save_result['remaining_output_path'],
            'remaining_filename': save_result['remaining_filename'],
            'weekly_rid_folder': save_result['weekly_rid_folder']  # FIXED: Added this
        }
        
    except Exception as e:
        print(f"ERROR in process_file_with_domain_subdomain_logic: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }



def calculate_subdomain_allocations(domain_config):
    """Calculate word allocations for each subdomain with remainder distribution"""
    allocation_plan = {}
    
    for domain_name, domain_data in domain_config.items():
        domain_total = domain_data['total']
        subdomains = domain_data['subdomains']
        
        if domain_total == 0:
            allocation_plan[domain_name] = {}
            continue
        
        # Calculate base allocation for each subdomain
        subdomain_allocations = []
        total_allocated = 0
        
        for subdomain_key, percentage in subdomains.items():
            base_words = int(domain_total * percentage / 100)
            subdomain_allocations.append({
                'key': subdomain_key,
                'percentage': percentage,
                'words': base_words
            })
            total_allocated += base_words
        
        # Calculate remainder and distribute to highest percentage subdomains
        remainder = domain_total - total_allocated
        
        if remainder > 0:
            # Sort by percentage (descending) for remainder distribution
            subdomain_allocations.sort(key=lambda x: x['percentage'], reverse=True)
            
            # Distribute remainder to highest percentage subdomains
            for i in range(remainder):
                if i < len(subdomain_allocations):
                    subdomain_allocations[i]['words'] += 1
        
        # Create final allocation mapping
        domain_allocation = {}
        for item in subdomain_allocations:
            domain_allocation[item['key']] = item['words']
        
        allocation_plan[domain_name] = domain_allocation
        
        print(f"{domain_name} allocation: {domain_allocation}")
    
    return allocation_plan


def extract_rows_by_domain_subdomain(df_original, allocation_plan):
    """Extract rows from dataframe according to domain/subdomain allocation plan"""
    try:
        extracted_rows = []
        extraction_summary = {}
        shortage_details = []
        has_shortages = False
        
        # Create a mapping of subdomain names to match against file data
        subdomain_mapping = {
            # HC subdomains
            'medical_devices': 'Medical Devices',
            'pharmaceutical': 'Pharmaceutical', 
            'biotechnology': 'Biotechnology',
            'healthcare_it': 'Healthcare IT',
            'clinical_diagnostic': 'Clinical Diagnostic',
            'medical_imaging': 'Medical Imaging',
            
            # CMFE subdomains
            'advanced_materials': 'Advanced Materials',
            'consumer_goods': 'Consumer Goods',
            'food_beverages': 'Food and Beverages',
            'bulk_chemicals': 'Bulk Chemicals',
            'specialty_chemicals': 'Specialty and Fine Chemicals',
            'energy': 'Energy',
            'packaging': 'Packaging',
            'food_ingredients': 'Food Ingredients',
            'polymers_resins': 'Polymers and Resins',
            'agrochemicals': 'Agrochemicals',
            'cosmetic_ingredients': 'Cosmetic Ingredients',
            'green_chemicals': 'Green Chemicals',
            
            # ICT subdomains
            'info_comm_tech': 'Information and Communication Technology',
            'automotive_transport': 'Automotive and Transportation',
            'industrial_automation': 'Industrial Automation and Machinery',
            'smart_technologies': 'Smart Technologies',
            'semiconductors': 'Semiconductors',
            'consumer_electronics': 'Consumer Electronics',
            'aerospace_defense': 'Aerospace and Defense',
            'construction_engineering': 'Construction Engineering'
        }
        
        # Track available rows for surplus redistribution
        available_surplus = {}
        
        # First pass: extract what's available for each domain/subdomain
        for domain_name, subdomain_allocations in allocation_plan.items():
            if not subdomain_allocations:  # Skip domains with 0 total
                continue
                
            extraction_summary[domain_name] = {
                'requested': sum(subdomain_allocations.values()),
                'extracted': 0,
                'subdomains': {}
            }
            
            for subdomain_key, requested_words in subdomain_allocations.items():
                if requested_words == 0:
                    continue
                    
                # Get actual subdomain name for matching
                actual_subdomain_name = subdomain_mapping.get(subdomain_key, subdomain_key)
                
                # Filter rows for this domain and subdomain
                domain_subdomain_rows = df_original[
                    (df_original['Domain'] == domain_name) & 
                    (df_original['Sub Domain'] == actual_subdomain_name)
                ].copy()
                
                available_count = len(domain_subdomain_rows)
                
                if available_count >= requested_words:
                    # Take exactly what was requested
                    selected_rows = domain_subdomain_rows.head(requested_words)
                    extracted_rows.append(selected_rows)
                    
                    # Track surplus for later redistribution
                    surplus = available_count - requested_words
                    if surplus > 0:
                        remaining_rows = domain_subdomain_rows.iloc[requested_words:]
                        available_surplus[f"{domain_name}_{subdomain_key}"] = remaining_rows
                    
                    extraction_summary[domain_name]['subdomains'][subdomain_key] = {
                        'requested': requested_words,
                        'extracted': requested_words,
                        'available': available_count,
                        'status': 'success'
                    }
                    extraction_summary[domain_name]['extracted'] += requested_words
                    
                else:
                    # Take all available rows (shortage situation)
                    if available_count > 0:
                        selected_rows = domain_subdomain_rows
                        extracted_rows.append(selected_rows)
                    
                    shortage = requested_words - available_count
                    has_shortages = True
                    shortage_details.append(
                        f"{domain_name}/{actual_subdomain_name}: Requested {requested_words}, "
                        f"only {available_count} available (shortage: {shortage})"
                    )
                    
                    extraction_summary[domain_name]['subdomains'][subdomain_key] = {
                        'requested': requested_words,
                        'extracted': available_count,
                        'available': available_count,
                        'shortage': shortage,
                        'status': 'shortage'
                    }
                    extraction_summary[domain_name]['extracted'] += available_count
        
        # Second pass: redistribute surplus to cover shortages
        if has_shortages and available_surplus:
            print(f"Redistributing surplus to cover shortages...")
            
            # Calculate total shortage needed
            total_shortage_needed = 0
            for domain_name, domain_summary in extraction_summary.items():
                for subdomain_key, subdomain_summary in domain_summary['subdomains'].items():
                    if subdomain_summary['status'] == 'shortage':
                        total_shortage_needed += subdomain_summary['shortage']
            
            # Use surplus to fill shortages
            surplus_used = 0
            for surplus_key, surplus_rows in available_surplus.items():
                if surplus_used >= total_shortage_needed:
                    break
                    
                rows_to_use = min(len(surplus_rows), total_shortage_needed - surplus_used)
                if rows_to_use > 0:
                    additional_rows = surplus_rows.head(rows_to_use)
                    extracted_rows.append(additional_rows)
                    surplus_used += rows_to_use
                    
                    shortage_details.append(
                        f"Used {rows_to_use} surplus words from {surplus_key.replace('_', '/')}"
                    )
        
        # Combine all extracted rows
        if extracted_rows:
            extracted_df = pd.concat(extracted_rows, ignore_index=True)
        else:
            extracted_df = pd.DataFrame()
        
        # Create remaining dataframe (original minus extracted)
        if not extracted_df.empty:
            # Get RIDs of extracted rows
            extracted_rids = set(extracted_df.index.tolist()) if 'RID' not in extracted_df.columns else set(extracted_df['RID'].tolist())
            
            # Create remaining dataframe
            if 'RID' in df_original.columns:
                remaining_df = df_original[~df_original['RID'].isin(extracted_rids)].copy()
            else:
                remaining_df = df_original.drop(extracted_df.index).copy()
        else:
            remaining_df = df_original.copy()
        
        return {
            'success': True,
            'extracted_df': extracted_df,
            'remaining_df': remaining_df,
            'extraction_summary': extraction_summary,
            'shortage_info': {
                'has_shortages': has_shortages,
                'shortage_details': shortage_details
            }
        }
        
    except Exception as e:
        print(f"Error in extract_rows_by_domain_subdomain: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def save_domain_extraction_results(extracted_df, remaining_df):
    """Save extraction results to files - FIXED VERSION"""
    try:
        today = datetime.today()
        
        # Save extracted rows as ROB.xlsx to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        rob_output_path = os.path.join(rpa_folder, "ROB.xlsx")
        extracted_df.to_excel(rob_output_path, index=False)
        
        # Save remaining rows with timestamp in date-wise folder
        weekly_rid_folder = os.path.join(
            r"C:\Users\vishwas\Desktop\RPA\Weekly_RID", 
            str(today.year), 
            f"{today.month:02d}", 
            f"{today.day:02d}"
        )
        os.makedirs(weekly_rid_folder, exist_ok=True)
        
        timestamp = f"{today.year}_{today.month:02d}_{today.day:02d}"
        remaining_filename = f"cleaned_weekly_rid_{timestamp}.xlsx"
        remaining_output_path = os.path.join(weekly_rid_folder, remaining_filename)
        remaining_df.to_excel(remaining_output_path, index=False)
        
        print(f"DEBUG: Files saved successfully")
        print(f"  ROB.xlsx: {rob_output_path}")
        print(f"  Remaining: {remaining_output_path}")
        
        return {
            'success': True,
            'rob_output_path': rob_output_path,
            'remaining_output_path': remaining_output_path,
            'remaining_filename': remaining_filename,
            'weekly_rid_folder': weekly_rid_folder  # FIXED: Added this
        }
        
    except Exception as e:
        print(f"ERROR in save_domain_extraction_results: {e}")
        return {
            'success': False,
            'error': str(e)
        }




@app.route('/download_remaining_rob')
def download_remaining_rob():
    """Download the remaining ROB file (original minus extracted rows)"""
    try:
        remaining_file_path = session.get('remaining_file_path')
        remaining_filename = session.get('remaining_filename', 'cleaned_rob_remaining.xlsx')
        
        if remaining_file_path and os.path.exists(remaining_file_path):
            return send_file(remaining_file_path, as_attachment=True, download_name=remaining_filename)
        else:
            flash('❌ Remaining ROB file not found. Please process a file first.')
            return redirect(url_for('rob'))
    except Exception as e:
        flash(f'❌ Error downloading remaining file: {str(e)}')
        return redirect(url_for('rob'))

@app.route('/download_extracted_rob')
def download_extracted_rob():
    """Download the extracted ROB.xlsx file and trigger OpenAI content generation"""
    try:
        rob_file_path = session.get('rob_file_path')
        
        if rob_file_path and os.path.exists(rob_file_path):
            
            # Start OpenAI content generation in background thread with 5-second delay
            print("🔄 Starting background OpenAI content generation...")
            threading.Thread(target=delayed_openai_content_generation, args=(5,)).start()
            
            return send_file(rob_file_path, as_attachment=True, download_name='ROB.xlsx')
        else:
            flash('❌ ROB.xlsx file not found. Please process a file first.')
            return redirect(url_for('rob'))
            
    except Exception as e:
        flash(f'❌ Error downloading ROB file: {str(e)}')
        return redirect(url_for('rob'))


def delayed_openai_content_generation(delay_seconds=5):
    """Generate content using OpenAI after delay (replaces Power Automate)"""
    try:
        print(f"⏳ Waiting {delay_seconds} seconds before starting content generation...")
        time.sleep(delay_seconds)
        
        print("🤖 Starting OpenAI content generation...")
        
        # ROB file path
        rob_file_path = r"C:\Users\vishwas\Desktop\RPA\ROB.xlsx"
        
        if not os.path.exists(rob_file_path):
            print("❌ ROB.xlsx file not found!")
            return
        
        # Check if API key is configured
        if OPENAI_CONFIG['API_KEY'] == 'your-openai-api-key-here':
            print("❌ OpenAI API key not configured! Using fallback content.")
        
        # Read ROB file
        df = pd.read_excel(rob_file_path)
        print(f"📊 Found {len(df)} markets in ROB file")
        
        # Create output directory with current date
        today = datetime.today()
        output_dir = os.path.join(
            r"C:\Users\vishwas\Desktop\RPA\Files",
            str(today.year),
            f"{today.month:02d}",
            f"{today.day:02d}"
        )
        os.makedirs(output_dir, exist_ok=True)
        print(f"📁 Output directory: {output_dir}")
        
        successful = 0
        failed = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                # Extract market data
                
                #print(f"[{index+1}/{len(df)}] Processing: {data['market_name']}")
                
                # Generate content
                content = generate_blog_from_row(row)
                
                # Save document
                success, filepath = save_market_document(row.get('Market Name'), content, output_dir)
                
                if success:
                    print(f"✅ Generated: {os.path.basename(filepath)}")
                    successful += 1
                else:
                    print(f"❌ Failed to save: (row.get('Market Name')")
                    failed += 1
                
                # Rate limiting
                time.sleep(2)
                
            except Exception as e:
                print(f"❌ Error processing  {e}")
                failed += 1
                continue
        
        print(f"✅ Content generation completed!")
        print(f"📊 Successful: {successful}, Failed: {failed}")
        print(f"📁 Files saved in: {output_dir}")
        
    except Exception as e:
        print(f"❌ Error in content generation: {e}")


# ============================================================================
# ADD THESE SIMPLE HELPER FUNCTIONS
# ============================================================================

import re
import openai

# OpenAI config
OPENAI_CONFIG = {
    'API_KEY': 'sk-proj-Vl99--DmgsXi7l1ivu8oKC8hmo4pDIBblicYkX_cHll6bEA',
    'MODEL': 'gpt-4.1-mini',  # Or 'gpt-4o' or 'gpt-3.5-turbo'
    'TEMPERATURE': 0.7
}

def generate_blog_from_row(row):
    """Extract data from row, format prompt, and generate blog via OpenAI."""
    try:
        # Extract values
        market_name = row.get('Market Name', 'Unknown Market')
        forecast_period = row.get('Forecast Period', 'Not specified')
        market_size_year = row.get('Market Size Year', 'Not specified')
        market_size = row.get('Market Size', '')
        cagr = row.get('CAGR', 'Not specified')
        key_players = row.get('Key Players', 'Not specified')

        # Extract size values from combined string
        match = re.search(r'USD ([\d.]+ Bn); Market Size in 2032: USD ([\d.]+ Bn)', market_size)
        if match:
            market_size_2025 = f"USD {match.group(1)}"
            market_size_2032 = f"USD {match.group(2)}"
        else:
            market_size_2025 = ""
            market_size_2032 = ""

        # Prepare prompt string from extracted values
        output = f"""
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➔ add first CTA link here, ➤Actionable Insights, ➤Market Segment and Regional coverage, ➔ add Second CTA link here, ➤Key players, ➤Growth factors, ➔add Third CTA link here, ➤ Market trends, ➤Market Opportunities, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet foe above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (first CTA link, Actionable Insights, Market Segment and Regional coverage, Second CTA link, Key players, Growth factors, Third CTA link, Market trends, Market Opportunities, and Frequently Asked Questions), this will increase the readibility. Cover content in in bullet pointers whenever possible each paragraph should be short, about 100 to 120 words. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should reflect exactly same in output as provided in input). Then First CTA link. Then Actionable Insights: In Actionable Insights, cover 3 to 4 actionable insights in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each actionable insights must have two sentence stats or actual instance examples from the recent year to support each point given in actionable insights, so that each given point look complete and meaningful. Next part is Market segment and regional Coverage where enlist the all subsegment under each segment categories and fragment region into given format. Comprehensive Segmentation and Classification of the Report: » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: » North America: U.S. and Canada » Latin America: Brazil, Argentina, Mexico, and Rest of Latin America » Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe » Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific » Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Second CTA link. Then key Players: List 12 to 20 highly relevant key players for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few key players, mentioning actual strategies and entities involved along with the actual outcome. Growth Factors: Growth factor heading and short description with supporting stats or examples from the recent year in the content. Then Add Third CTA link. Then Market Trends: Market Trend heading and short description with supporting stats or examples from the recent year in the content. Then Market Opportunities: Provide 3 to 4 market opportunities, 2-3 opportunities based upon segment and one opportunity based upon region. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market. Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team 

Market Name and Data:
Market Name- {market_name}
CAGR: {cagr}
Forecast period is: {forecast_period}
Market Size for {market_size_year} is {market_size_2025}
Market Size in 2032: {market_size_2032}
Key players: {key_players}

Key Reasons for Buying the (insert market name here) Report: ✦ Comprehensive analysis of the changing competitive landscape ✦ Assists in decision-making processes for the businesses along with detailed strategic planning methodologies ✦ The report offers forecast data and an assessment of the (insert market name here) ✦ Helps in understanding the key product segments and their estimated growth rate ✦ In-depth analysis of market drivers, restraints, trends, and opportunities ✦ Comprehensive regional analysis of the (insert market name here) ✦ Extensive profiling of the key stakeholders of the business sphere ✦ Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, market trends, market opportunities, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and Market Opportunities where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ➤Actionable Insights, ➤Market Segment and Regional Coverage, ➔ Inserted Second CTA link, ➤Key Players, ➤Growth factors, ➔ Inserted Third CTA link, ➤ Market Trends, ➤Market Opportunities, and ❓ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Actionable Insights ●, Market Segment and Regional Coverage● , Key players●, Growth Factors●,  Market Trends●, Market Opportunities●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
"""

        # Send to OpenAI
        client = openai.OpenAI(api_key=OPENAI_CONFIG['API_KEY'])

        response = client.chat.completions.create(
            model=OPENAI_CONFIG['MODEL'],
            messages=[
                {"role": "user", "content": output}
            ],
            temperature=OPENAI_CONFIG.get('TEMPERATURE', 0.7)
        )

        return response.choices[0].message.content.strip()

    except Exception as e:
        print(f"OpenAI error: {e}")
        return "Error generating content."

def save_market_document(market_name, content, output_folder):
    """Save content as Word document"""
    try:
        doc = Document()
        doc.add_heading(f"{market_name} - Market Research Report", level=1)
        
        # Add content paragraphs
        for paragraph in content.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Add contact info
      
        
        # Save file
        safe_name = "".join(c for c in market_name if c.isalnum() or c in (' ', '_')).strip()
        filename = f"ROB_{safe_name}.doc"
        filepath = os.path.join(output_folder, filename)
        doc.save(filepath)
        
        return True, filepath
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, None


'''app.route('/api/auto_trigger_power_automate', methods=['POST'])
def auto_trigger_power_automate():
    """API endpoint for auto-triggering Power Automate"""
    try:
        # Check if we should trigger (based on recent download)
        if session.get('trigger_power_automate'):
            # Clear the flag
            session['trigger_power_automate'] = False
            
            # Trigger in background
            threading.Thread(target=delayed_power_automate_trigger, args=(0,)).start()
            
            return jsonify({
                'status': 'success', 
                'message': 'Power Automate triggered automatically after ROB download'
            })
        else:
            return jsonify({
                'status': 'error', 
                'message': 'No recent ROB download detected'
            })
    except Exception as e:
        return jsonify({
            'status': 'error', 
            'message': f'Error: {str(e)}'
        })'''

# ============================================================================
# WEEKLY REPORT ROUTES
# ============================================================================
# ============================================================================
# UPDATED WEEKLY REPORT ROUTES WITH DOMAIN DISTRIBUTION
# ============================================================================
# ============================================================================
# HELPER FUNCTIONS FOR RID CLEANING
# ============================================================================

def clean_rid(rid):
    """Convert RID to clean integer string, handling floats like 56.0 -> 56"""
    try:
        # Handle float values like 56.0 -> 56
        if isinstance(rid, (int, float)) and not pd.isna(rid):
            return str(int(float(rid)))
        # Handle string values
        elif isinstance(rid, str):
            rid = rid.strip()
            if rid and rid.lower() != 'nan':
                # Try to clean float strings like '56.0' -> '56'
                try:
                    return str(int(float(rid)))
                except ValueError:
                    return rid
        return str(rid)
    except (ValueError, TypeError):
        return str(rid).strip()


def clean_rid_list(rid_list):
    """Clean RID list to handle float values consistently"""
    return [clean_rid(rid) for rid in rid_list]


def test_rid_cleaning():
    """Test function to demonstrate RID cleaning"""
    test_rids = [56.0, 20.0, '78.0', '99', 45, 'ABC123', None]
    cleaned = clean_rid_list(test_rids)
    print("RID Cleaning Test:")
    for original, cleaned_val in zip(test_rids, cleaned):
        print(f"  {original} -> {cleaned_val}")
    return cleaned


# ============================================================================
# UPDATED WEEKLY REPORT ROUTES WITH DOMAIN DISTRIBUTION
# ============================================================================

# Update the main weekly_report route to use improved workflow
# ============================================================================
# WEEKLY REPORT ROUTES - CLEAN & CORRECTED VERSION
# ============================================================================

@app.route('/weekly_report', methods=['GET', 'POST'])
def weekly_report():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        if form_type == 'backend_processing':
            return handle_backend_processing()
        else:
            return handle_rid_analysis_with_domains()
    
    # GET request - show form
    return render_template('weekly_report.html', qualified_rids=None, filter_summary=None, backend_result=None)


def handle_rid_analysis_with_domains():
    """Handle RID analysis with improved domain distribution workflow"""
    try:
        print("Improved RID Analysis with Domain Distribution POST request received!")
        
        # Get filter parameters from form
        min_search_volume = int(request.form.get('min_search_volume', 5000))
        competition_level = request.form.get('competition_level', 'Low')
        analyze_trends = request.form.get('analyze_trends') == 'on'
        target_count = request.form.get('target_count')
        
        # Convert target_count to int if provided
        if target_count and target_count.strip():
            try:
                target_count = int(target_count)
                if target_count <= 0:
                    raise ValueError("Target count must be positive")
            except ValueError:
                flash('❌ Please enter a valid target word count!')
                return redirect(request.url)
        else:
            target_count = None
        
        print(f"User Filters: Search >= {min_search_volume}, Competition = {competition_level}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        print(f"Target Count: {target_count if target_count else 'Not specified'}")
        
        # Validate form inputs
        if not min_search_volume or min_search_volume < 0:
            flash('❌ Please enter a valid minimum search volume!')
            return redirect(request.url)
            
        if not competition_level:
            flash('❌ Please select a competition level!')
            return redirect(request.url)
        
        # Handle file uploads
        ranking_file = request.files.get('ranking_file')
        if not ranking_file or ranking_file.filename == '':
            flash('❌ Please select a ranking Excel file!')
            return redirect(request.url)

        if not allowed_file(ranking_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
            return redirect(request.url)

        rob_file = request.files.get('cleaned_rob_file')
        if not rob_file or rob_file.filename == '':
            flash('❌ Please select a cleaned ROB Excel file!')
            return redirect(request.url)

        if not allowed_file(rob_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ROB file!')
            return redirect(request.url)

        # Save uploaded files
        ranking_filename = secure_filename(ranking_file.filename)
        ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
        ranking_file.save(ranking_path)
        
        rob_filename = secure_filename(rob_file.filename)
        rob_path = os.path.join(app.config['UPLOAD_FOLDER'], rob_filename)
        rob_file.save(rob_path)
        
        # Process files with improved workflow
        result_summary = process_dual_files_improved_workflow(
            ranking_path, rob_path, min_search_volume, competition_level, analyze_trends, target_count
        )
        
        # Display comprehensive results
        if result_summary['success']:
            # Main success message
            flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
            
            # Show original counts
            original_stats = result_summary.get("original_stats", {})
            if original_stats:
                flash(f'📊 Original Data: Total={original_stats.get("total", 0)}, HC={original_stats.get("HC", 0)}, ICT={original_stats.get("ICT", 0)}, CMFE={original_stats.get("CMFE", 0)}')
            
            # Show target vs achieved
            target_info = result_summary.get("target_info", {})
            if target_info:
                flash(f'🎯 Target: {target_info.get("target_total", 0)} words ({target_info.get("target_per_domain", 0)} per domain)')
                flash(f'🎯 Achieved: HC={target_info.get("achieved_HC", 0)}, ICT={target_info.get("achieved_ICT", 0)}, CMFE={target_info.get("achieved_CMFE", 0)}')
            
            # Show shortage information
            shortage_info = result_summary.get("shortage_info", {})
            if shortage_info.get("has_shortage"):
                for domain, shortage in shortage_info.get("shortages", {}).items():
                    if shortage > 0:
                        flash(f'⚠️ {domain} Domain Shortage: {shortage} words (filled from other domains)')
                
                total_shortage = shortage_info.get("total_shortage", 0)
                if total_shortage > 0:
                    flash(f'💡 Recommendation: Add at least {total_shortage} more {", ".join(shortage_info.get("short_domains", []))} domain words next time')
            
            # Show search volume adjustments
            search_adjustments = result_summary.get("filter_summary", {}).get("search_adjustments", {})
            if search_adjustments.get("adjusted"):
                original_sv = search_adjustments.get("original_volume", min_search_volume)
                final_sv = search_adjustments.get("final_volume", min_search_volume)
                attempts = search_adjustments.get("attempts", [])
                
                if len(attempts) > 1:
                    if final_sv == 0:
                        flash(f'🔄 Auto-adjusted search volume: {original_sv:,} → No Filter (tried: {", ".join([f"{v:,}" if v > 0 else "No Filter" for v in attempts])})')
                    else:
                        flash(f'🔄 Auto-adjusted search volume: {original_sv:,} → {final_sv:,} (tried: {", ".join([f"{v:,}" if v > 0 else "No Filter" for v in attempts])})')
                else:
                    if final_sv == 0:
                        flash(f'📈 Used search volume: No Filter (all keywords)')
                    else:
                        flash(f'📈 Used search volume: {final_sv:,}')
            else:
                flash(f'📈 Used search volume: {min_search_volume:,}')
           
           # Show Google Trends progressive results
            trends_info = result_summary.get("trends_info", {})
            if trends_info.get("enabled") and trends_info.get("breakdown"):
                breakdown = trends_info["breakdown"] 
                flash(f'🔍 Google Trends: {breakdown["medium_2plus"]} (2+), {breakdown["relaxed_1plus"]} (1+) = {breakdown["total_qualified"]} total')
            # Show final results vs target
            if target_count:
                if result_summary["qualified_rids_count"] < target_count:
                    final_shortage = target_count - result_summary["qualified_rids_count"]
                    flash(f'⚠️ Only {result_summary["qualified_rids_count"]} words qualifying out of {target_count} requested')
                    flash(f'💡 Please add {final_shortage} more words to your input file for next time')
                elif result_summary["qualified_rids_count"] == target_count:
                    flash(f'🎯 Perfect! Achieved exactly your target of {target_count} words')
                else:
                    flash(f'✨ Exceeded target! Got {result_summary["qualified_rids_count"]} out of {target_count} requested')
            
            # ROB extraction results
            flash(f'✅ Extracted {result_summary["matched_rob_rows"]} matching ROB rows')
            flash(f'📁 Weekly ROB.xlsx saved to Desktop/RPA folder!')
            
            # Updated ranking sheet info
            if result_summary.get("updated_ranking_path"):
                flash(f'📄 Updated ranking sheet saved and ready for download!')
            
        else:
            flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
            result_summary = None
        
        # Clean up uploaded files
        try:
            os.remove(ranking_path)
            os.remove(rob_path)
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up files: {cleanup_error}")
        
        # Render template with results
        return render_template('weekly_report.html', 
                              qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                              filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                              backend_result=None,
                              rob_extraction_result=result_summary)
        
    except Exception as e:
        print(f"Error: {e}")
        flash(f'❌ Error processing files: {str(e)}')
        return redirect(request.url)


# ============================================================================
# MAIN PROCESSING FUNCTION - SINGLE CLEAN VERSION
# ============================================================================

def process_dual_files_improved_workflow(ranking_path, rob_path, min_search_volume, competition_level, analyze_trends, target_count=None):
    """Process files with improved workflow: Domain Distribution → Trends → Search → Competition → RID Removal → ROB Matching"""
    try:
        from flask import session
        
        print(f"\n=== IMPROVED WORKFLOW: DOMAIN DISTRIBUTION FIRST ===")
        print(f"Ranking file: {ranking_path}")
        print(f"ROB file: {rob_path}")
        print(f"Target count: {target_count if target_count else 'Not specified'}")
        
        # STEP 1: Load and analyze original data
        print("\n📊 STEP 1: Analyzing original data...")
        try:
            if ranking_path.endswith('.csv'):
                df_original = pd.read_csv(ranking_path)
            else:
                # Try different engines for reading Excel files
                try:
                    df_original = pd.read_excel(ranking_path, engine='openpyxl')
                except Exception as e1:
                    print(f"Failed with openpyxl, trying xlrd: {e1}")
                    try:
                        df_original = pd.read_excel(ranking_path, engine='xlrd')
                    except Exception as e2:
                        print(f"Failed with xlrd, trying default engine: {e2}")
                        df_original = pd.read_excel(ranking_path)
        except Exception as e:
            return {'success': False, 'error': f"Could not read ranking file: {str(e)}"}
        
        # Validate required columns
        required_columns = ['AVG. Search', 'Competition', 'RID', 'Domain']
        missing_columns = [col for col in required_columns if col not in df_original.columns]
        if missing_columns:
            return {'success': False, 'error': f"Missing required columns: {missing_columns}"}
        
        # Get original statistics
        original_total = len(df_original)
        original_domain_counts = df_original['Domain'].value_counts().to_dict()
        
        print(f"Original Data: Total={original_total}")
        print(f"Original Domain Counts: {original_domain_counts}")
        
        original_stats = {
            'total': original_total,
            'HC': original_domain_counts.get('HC', 0),
            'ICT': original_domain_counts.get('ICT', 0),
            'CMFE': original_domain_counts.get('CMFE', 0)
        }
        
        # STEP 2: Apply improved domain distribution
        print(f"\n🎯 STEP 2: Applying equal domain distribution...")
        distributed_df, distribution_result = apply_improved_domain_distribution(df_original, target_count)
        
        if distributed_df.empty:
            return {
                'success': False,
                'error': 'No data after domain distribution'
            }
        
        print(f"After domain distribution: {len(distributed_df)} rows")
        print(f"Domain distribution: {get_domain_counts(distributed_df)}")
        
        # STEP 3: Apply Google Trends filter (if enabled) with progressive filtering
        trends_filtered_df = distributed_df.copy()
        trends_info = {'enabled': analyze_trends, 'message': 'Google Trends disabled'}
        
        if analyze_trends:
            print(f"\n🔍 STEP 3: Applying Progressive Google Trends filter...")
            
            if not GOOGLE_TRENDS_CONFIG.get('API_KEY') or GOOGLE_TRENDS_CONFIG['API_KEY'] == 'YOUR_API_KEY_HERE':
                trends_info['message'] = "Google Trends disabled - no API key configured"
                print("⚠️ No API key configured - skipping Google Trends")
            else:
                keywords_data = distributed_df.to_dict('records')
                trending_data = analyze_keywords_with_progressive_trends(keywords_data, target_count)
                trending_rids = [item['RID'] for item in trending_data if 'RID' in item]
                
                if trending_rids:
                    trends_filtered_df = distributed_df[distributed_df['RID'].isin(trending_rids)].copy()
                    trends_count = len(trends_filtered_df)
                    
                    # Get breakdown info
                    breakdown = {
                        'strict_3plus': len([x for x in trending_data if x.get('qualified_at') == '3+']),
                        'medium_2plus': len([x for x in trending_data if x.get('qualified_at') == '2+']),
                        'relaxed_1plus': len([x for x in trending_data if x.get('qualified_at') == '1+']),
                        'total_qualified': len(trending_data)
                    }
                    
                    trends_info['message'] = f"Progressive Google Trends: {trends_count} qualified keywords"
                    trends_info['breakdown'] = breakdown
                    print(f"After Progressive Google Trends: {trends_count} rows")
                else:
                    trends_filtered_df = pd.DataFrame()
                    trends_info['message'] = "Google Trends: No trending keywords found"
                    trends_info['breakdown'] = {'strict_3plus': 0, 'medium_2plus': 0, 'relaxed_1plus': 0, 'total_qualified': 0}
                    print("No trending keywords found")
        
        # STEP 4: Apply adaptive search volume filter
        print(f"\n📈 STEP 4: Applying adaptive search volume filter...")
        search_filtered_df, final_search_volume, search_adjustments = apply_adaptive_search_volume_filter(
            trends_filtered_df, min_search_volume, target_count
        )
        search_count = len(search_filtered_df)
        
        # STEP 5: Apply competition filter
        print(f"\n🏆 STEP 5: Applying competition filter ({competition_level})...")
        if len(search_filtered_df) > 0:
            final_filtered_df = apply_competition_filter_with_priority(search_filtered_df, competition_level)
            final_count = len(final_filtered_df)
            print(f"After competition filter: {final_count} rows")
        else:
            final_filtered_df = pd.DataFrame()
            final_count = 0
        
        # Get final qualified RIDs
        qualified_rids = clean_rid_list(final_filtered_df['RID'].tolist()) if not final_filtered_df.empty else []
        
        # STEP 6: Remove qualified RIDs from ranking sheet
        print(f"\n📝 STEP 6: Removing qualified RIDs from original ranking sheet...")
        updated_ranking_path = None
        if qualified_rids:
            updated_ranking_path = remove_trending_rids_from_ranking(df_original, qualified_rids)
            if updated_ranking_path:
                session['updated_ranking_path'] = updated_ranking_path
                print(f"✅ Updated ranking sheet saved and ready for download")
            else:
                print(f"❌ Failed to save updated ranking sheet")
        else:
            print(f"⚠️ No qualified RIDs to remove")
        
        # STEP 7: Process ROB file
        print(f"\n📋 STEP 7: Processing ROB file...")
        if not qualified_rids:
            return {
                'success': False,
                'error': 'No qualified RIDs found after applying all filters'
            }
        
        # Read ROB file and extract matching rows - FIXED VERSION
        try:
            if rob_path.endswith('.csv'):
                rob_df = pd.read_csv(rob_path)
            else:
                # Try different engines for reading Excel files
                try:
                    rob_df = pd.read_excel(rob_path, engine='openpyxl')
                    print(f"✅ Successfully read ROB file with openpyxl engine")
                except Exception as e1:
                    print(f"Failed with openpyxl, trying xlrd: {e1}")
                    try:
                        rob_df = pd.read_excel(rob_path, engine='xlrd')
                        print(f"✅ Successfully read ROB file with xlrd engine")
                    except Exception as e2:
                        print(f"Failed with xlrd, trying default engine: {e2}")
                        rob_df = pd.read_excel(rob_path)
                        print(f"✅ Successfully read ROB file with default engine")
        except Exception as e:
            return {
                'success': False,
                'error': f'Could not read ROB file: {str(e)}'
            }
        
        print(f"ROB file loaded successfully. Shape: {rob_df.shape}")
        print(f"ROB file columns: {list(rob_df.columns)}")
        
        # Find Report ID column
        report_id_column = None
        possible_columns = ['Report ID', 'ReportID', 'report_id', 'ID', 'Report_ID', 'Market Name', 'RID']
        for col in possible_columns:
            if col in rob_df.columns:
                report_id_column = col
                print(f"Found Report ID column: '{col}'")
                break
        
        if not report_id_column:
            return {
                'success': False,
                'error': f'Report ID column not found in ROB file. Available columns: {list(rob_df.columns)}'
            }
        
        # Clean and match ROB data
        print(f"Cleaning and matching ROB data using column: '{report_id_column}'")
        qualified_rids_clean = clean_rid_list(qualified_rids)
        rob_df[report_id_column] = rob_df[report_id_column].apply(clean_rid)
        
        print(f"Qualified RIDs (sample): {qualified_rids_clean[:5]}")
        print(f"ROB RIDs (sample): {rob_df[report_id_column].head().tolist()}")
        
        matching_rob_rows = rob_df[rob_df[report_id_column].isin(qualified_rids_clean)].copy()
        print(f"Found {len(matching_rob_rows)} matching ROB rows")
        
        # Save results
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        output_path = os.path.join(rpa_folder, "weekly_RID.xlsx")
        
        # Save with proper error handling
        try:
            with pd.ExcelWriter(output_path, engine='xlsxwriter') as writer:
                matching_rob_rows.to_excel(writer, index=False, sheet_name='ROB_Data')
            print(f"✅ weekly_RID.xlsx saved to: {output_path}")
        except Exception as e:
            print(f"Error saving Excel file: {e}")
            # Fallback to openpyxl for saving
            try:
                with pd.ExcelWriter(output_path, engine='openpyxl') as writer:
                    matching_rob_rows.to_excel(writer, index=False, sheet_name='ROB_Data')
                print(f"✅ weekly_RID.xlsx saved to: {output_path} (using openpyxl)")
            except Exception as e2:
                return {
                    'success': False,
                    'error': f'Could not save output file: {str(e2)}'
                }
        
        # Create comprehensive result summary
        final_domain_counts = get_domain_counts(final_filtered_df)
        
        result_summary = {
            'success': True,
            'qualified_rids': qualified_rids,
            'qualified_rids_count': len(qualified_rids),
            'matched_rob_rows': len(matching_rob_rows),
            'output_path': output_path,
            'updated_ranking_path': updated_ranking_path,
            'original_stats': original_stats,
            'target_info': {
                'target_total': target_count if target_count else len(qualified_rids),
                'target_per_domain': (target_count // 3) if target_count else (len(qualified_rids) // 3),
                'achieved_HC': final_domain_counts.get('HC', 0),
                'achieved_ICT': final_domain_counts.get('ICT', 0),
                'achieved_CMFE': final_domain_counts.get('CMFE', 0)
            },
            'shortage_info': distribution_result.get('shortage_info', {}),
            'trends_info': trends_info,
            'filter_summary': {
                'original_count': original_total,
                'distributed_count': len(distributed_df),
                'trends_count': len(trends_filtered_df),
                'search_filtered_count': search_count,
                'final_count': final_count,
                'min_search': f"{final_search_volume:,}" if final_search_volume > 0 else "No Filter",
                'original_search_volume': f"{min_search_volume:,}",
                'search_adjustments': search_adjustments,
                'competition': competition_level
            }
        }
        
        return result_summary
        
    except Exception as e:
        print(f"Error in improved workflow: {e}")
        import traceback
        traceback.print_exc()
        return {
            'success': False,
            'error': str(e)
        }

# ============================================================================
# HELPER FUNCTIONS
# ============================================================================

def apply_improved_domain_distribution(df, target_count):
    """Apply improved domain distribution with equal proportions as close as possible"""
    try:
        total_available = len(df)
        domains = ['HC', 'ICT', 'CMFE']
        
        # Use all data if no target specified
        if not target_count:
            target_count = total_available
        
        # Calculate target per domain (as close to equal as possible)
        target_per_domain = target_count // 3
        remainder = target_count % 3
        
        print(f"Total available: {total_available}")
        print(f"Target total: {target_count}")
        print(f"Target per domain: {target_per_domain} (+ {remainder} remainder)")
        
        # Get domain counts
        domain_counts = df['Domain'].value_counts()
        print(f"Available by domain: {domain_counts.to_dict()}")
        
        selected_rows = []
        shortage_info = {
            'has_shortage': False,
            'shortages': {},
            'total_shortage': 0,
            'short_domains': [],
            'surplus_used': {}
        }
        
        # PHASE 1: Try to get target_per_domain from each domain
        remaining_needed = 0
        domain_targets = {}
        
        for i, domain in enumerate(domains):
            # Add remainder to first few domains
            current_target = target_per_domain + (1 if i < remainder else 0)
            domain_targets[domain] = current_target
            
            domain_df = df[df['Domain'] == domain].copy()
            available = len(domain_df)
            
            # Sort by priority (Low competition first, then high search volume)
            domain_df = sort_by_priority(domain_df)
            
            if available >= current_target:
                # Take exactly what we need
                selected = domain_df.head(current_target)
                selected_rows.append(selected)
                print(f"{domain}: Selected {current_target} out of {available} available")
            else:
                # Take all available and note shortage
                selected = domain_df
                selected_rows.append(selected)
                shortage = current_target - available
                remaining_needed += shortage
                
                shortage_info['has_shortage'] = True
                shortage_info['shortages'][domain] = shortage
                shortage_info['total_shortage'] += shortage
                shortage_info['short_domains'].append(domain)
                
                print(f"{domain}: Selected {available} out of {current_target} needed (shortage: {shortage})")
        
        # Combine initially selected rows
        result_df = pd.concat(selected_rows, ignore_index=True) if selected_rows else pd.DataFrame()
        selected_rids = set(result_df['RID'].tolist()) if not result_df.empty else set()
        
        # PHASE 2: Fill shortages from domains with surplus
        if remaining_needed > 0:
            print(f"\n🔄 FILLING SHORTAGES: Need {remaining_needed} more rows")
            
            for domain in domains:
                if remaining_needed <= 0:
                    break
                
                # Get unselected rows from this domain
                domain_df = df[df['Domain'] == domain].copy()
                unselected_domain_df = domain_df[~domain_df['RID'].isin(selected_rids)]
                unselected_domain_df = sort_by_priority(unselected_domain_df)
                
                available_surplus = len(unselected_domain_df)
                take_count = min(remaining_needed, available_surplus)
                
                if take_count > 0:
                    additional_rows = unselected_domain_df.head(take_count)
                    result_df = pd.concat([result_df, additional_rows], ignore_index=True)
                    selected_rids.update(additional_rows['RID'].tolist())
                    remaining_needed -= take_count
                    
                    shortage_info['surplus_used'][domain] = take_count
                    print(f"{domain} surplus: Used {take_count} rows (remaining needed: {remaining_needed})")
        
        # Final sorting by priority
        result_df = sort_by_priority(result_df)
        
        final_count = len(result_df)
        final_domain_counts = get_domain_counts(result_df)
        
        print(f"✅ FINAL DISTRIBUTED COUNT: {final_count}")
        print(f"📊 FINAL DOMAIN DISTRIBUTION: {final_domain_counts}")
        
        distribution_result = {
            'shortage_info': shortage_info,
            'target_per_domain': target_per_domain,
            'final_distribution': final_domain_counts
        }
        
        return result_df, distribution_result
        
    except Exception as e:
        print(f"Error in improved domain distribution: {e}")
        raise e


def get_domain_counts(df):
    """Get domain distribution counts"""
    if df.empty:
        return {'HC': 0, 'ICT': 0, 'CMFE': 0}
    
    domain_counts = df['Domain'].value_counts().to_dict()
    return {
        'HC': domain_counts.get('HC', 0),
        'ICT': domain_counts.get('ICT', 0),
        'CMFE': domain_counts.get('CMFE', 0)
    }


def sort_by_priority(df):
    """Sort dataframe by competition priority (Low first) and search volume (high first)"""
    if df.empty:
        return df
        
    # Create priority mapping
    priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
    df_copy = df.copy()
    df_copy['competition_priority'] = df_copy['Competition'].map(priority_map).fillna(4)
    
    # Sort by priority (Low first) and search volume (high first)
    df_sorted = df_copy.sort_values(['competition_priority', 'AVG. Search'], ascending=[True, False])
    
    # Remove the temporary column
    df_sorted = df_sorted.drop('competition_priority', axis=1)
    
    return df_sorted


def apply_competition_filter_with_priority(df, competition_level):
    """Apply competition filter with priority system"""
    if df.empty:
        return df
        
    if competition_level == 'All':
        return sort_by_priority(df)
    
    # Priority order: Low -> Medium -> High
    priority_order = ['Low', 'Medium', 'High']
    
    if competition_level in priority_order:
        selected_index = priority_order.index(competition_level)
        allowed_levels = priority_order[:selected_index + 1]  # Include all levels up to selected
        
        print(f"Competition priority: {' → '.join(allowed_levels)} (up to {competition_level})")
        
        # Filter for allowed competition levels
        filtered_df = df[df['Competition'].isin(allowed_levels)].copy()
        
        # Sort by priority order and search volume
        filtered_df = sort_by_priority(filtered_df)
        
        return filtered_df
    else:
        # Fallback to exact match
        return df[df['Competition'] == competition_level].copy()


def apply_adaptive_search_volume_filter(df, original_min_search, target_count=None):
    """Apply search volume filter with automatic adjustment: 5000 → 500 → 50 → 0"""
    try:
        if df.empty:
            return df, original_min_search, {'adjusted': False, 'attempts': []}
        
        # Define search volume thresholds to try in order (including 0 as final fallback)
        search_thresholds = [original_min_search,50000,5000,500, 50, 0]
        
        # Remove duplicates and sort in descending order
        search_thresholds = sorted(list(set(search_thresholds)), reverse=True)
        
        print(f"Search volume adjustment enabled. Will try: {search_thresholds}")
        
        search_adjustments = {
            'adjusted': False,
            'original_volume': original_min_search,
            'final_volume': original_min_search,
            'attempts': [],
            'results': {}
        }
        
        best_result = None
        final_search_volume = original_min_search
        
        for threshold in search_thresholds:
            print(f"  Trying search volume >= {threshold:,}...")
            
            # Apply current threshold
            if threshold == 0:
                # No search volume filter - take all data
                filtered_df = df.copy()
                print(f"    No search volume filter applied (threshold = 0)")
            else:
                filtered_df = df[df['AVG. Search'] >= threshold].copy()
            
            current_count = len(filtered_df)
            
            search_adjustments['attempts'].append(threshold)
            search_adjustments['results'][threshold] = current_count
            
            if threshold == 0:
                print(f"    Result: {current_count} rows with no search volume filter")
            else:
                print(f"    Result: {current_count} rows with search volume >= {threshold:,}")
            
            # Store this result
            best_result = filtered_df.copy()
            final_search_volume = threshold
            
            # Check if we have enough data
            if target_count:
                if current_count >= target_count:
                    if threshold == 0:
                        print(f"    ✅ Sufficient data found with no search volume filter (target: {target_count})")
                    else:
                        print(f"    ✅ Sufficient data found with {threshold:,} (target: {target_count})")
                    break
                else:
                    shortage = target_count - current_count
                    if threshold == 0:
                        print(f"    ⚠️ Even with no search volume filter, only {current_count} available (target: {target_count})")
                        print(f"    📝 Final result: {current_count} keywords (all available after other filters)")
                    else:
                        print(f"    ⚠️ Still short by {shortage} with {threshold:,}, trying lower threshold...")
            else:
                # If no target specified, use the first (highest) threshold
                if threshold == 0:
                    print(f"    ✅ Using no search volume filter (no target specified)")
                else:
                    print(f"    ✅ Using {threshold:,} (no target specified)")
                break
        
        # Update adjustment info
        if final_search_volume != original_min_search:
            search_adjustments['adjusted'] = True
            search_adjustments['final_volume'] = final_search_volume
            if final_search_volume == 0:
                print(f"🔄 Search volume auto-adjusted: {original_min_search:,} → No Filter (0)")
            else:
                print(f"🔄 Search volume auto-adjusted: {original_min_search:,} → {final_search_volume:,}")
        else:
            print(f"📈 Using original search volume: {original_min_search:,}")
        
        search_adjustments['final_volume'] = final_search_volume
        
        return best_result if best_result is not None else pd.DataFrame(), final_search_volume, search_adjustments
        
    except Exception as e:
        print(f"Error in adaptive search volume filter: {e}")
        # Fallback to original behavior
        filtered_df = df[df['AVG. Search'] >= original_min_search].copy() if not df.empty else pd.DataFrame()
        return filtered_df, original_min_search, {'adjusted': False, 'attempts': [original_min_search], 'error': str(e)}


def clean_rid_list(rid_list):
    """Clean RID list to handle float values consistently"""
    return [clean_rid(rid) for rid in rid_list]


def clean_rid(rid):
    """Convert RID to clean integer string, handling floats like 56.0 -> 56"""
    try:
        # Handle float values like 56.0 -> 56
        if isinstance(rid, (int, float)) and not pd.isna(rid):
            return str(int(float(rid)))
        # Handle string values
        elif isinstance(rid, str):
            rid = rid.strip()
            if rid and rid.lower() != 'nan':
                # Try to clean float strings like '56.0' -> '56'
                try:
                    return str(int(float(rid)))
                except ValueError:
                    return rid
        return str(rid)
    except (ValueError, TypeError):
        return str(rid).strip()


def remove_trending_rids_from_ranking(df_original, qualified_rids):
    """Remove qualified RIDs from original ranking sheet and save updated version"""
    try:
        if not qualified_rids:
            print("No qualified RIDs to remove")
            return None
            
        print(f"Original ranking sheet has {len(df_original)} rows")
        print(f"Removing {len(qualified_rids)} qualified RIDs...")
        
        # Clean RIDs to ensure proper matching
        qualified_rids_clean = clean_rid_list(qualified_rids)
        df_original_copy = df_original.copy()
        df_original_copy['RID'] = df_original_copy['RID'].apply(clean_rid)
        
        print(f"Sample RIDs to remove: {qualified_rids_clean[:5]}...")
        print(f"Sample RIDs in original: {df_original_copy['RID'].head().tolist()}")
        
        # Create updated dataframe without qualified RIDs
        df_updated = df_original_copy[~df_original_copy['RID'].isin(qualified_rids_clean)].copy()
        
        rows_removed = len(df_original_copy) - len(df_updated)
        print(f"Rows removed: {rows_removed}")
        print(f"Remaining rows: {len(df_updated)}")
        
        if rows_removed == 0:
            print("⚠️ Warning: No rows were actually removed. Check RID matching.")
        
        # Create output directory structure
        today = datetime.today()
        ranking_base_folder = os.path.join(r"C:\Users\vishwas\Desktop\RPA\Ranking sheet", str(today.year), f"{today.month:02d}")
        os.makedirs(ranking_base_folder, exist_ok=True)
        
        # Create filename with timestamp
        updated_filename = f"updated_ranking_sheet_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        updated_file_path = os.path.join(ranking_base_folder, updated_filename)
        
        # Save updated ranking sheet
        print(f"Saving updated ranking sheet to: {updated_file_path}")
        
        # Use xlsxwriter for better compatibility
        with pd.ExcelWriter(updated_file_path, engine='xlsxwriter') as writer:
            df_updated.to_excel(writer, index=False, sheet_name='Updated_Ranking')
        
        # Verify file was created
        if os.path.exists(updated_file_path):
            file_size = os.path.getsize(updated_file_path)
            print(f"✅ Updated ranking sheet saved successfully!")
            print(f"   File: {updated_file_path}")
            print(f"   Size: {file_size} bytes")
            print(f"   Rows: {len(df_updated)}")
            print(f"   Removed: {rows_removed} qualified RIDs")
            return updated_file_path
        else:
            print(f"❌ File was not created successfully")
            return None
        
    except Exception as e:
        print(f"Error removing qualified RIDs from ranking sheet: {e}")
        import traceback
        traceback.print_exc()
        return None


# ============================================================================
# DOWNLOAD ROUTES
# ============================================================================

@app.route('/download_updated_ranking')
def download_updated_ranking():
    """Download the updated ranking sheet (with qualified rows removed)"""
    try:
        from flask import session
        
        # Get file path from session
        updated_ranking_path = session.get('updated_ranking_path')
        
        if updated_ranking_path and os.path.exists(updated_ranking_path):
            # Get just the filename for download
            filename = os.path.basename(updated_ranking_path)
            return send_file(updated_ranking_path, as_attachment=True, download_name=filename)
        else:
            # Fallback: try to find today's file
            today = datetime.today()
            ranking_base_folder = os.path.join(r"C:\Users\vishwas\Desktop\RPA\Ranking sheet", str(today.year), f"{today.month:02d}")
            updated_filename = f"updated_ranking_sheet_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
            fallback_path = os.path.join(ranking_base_folder, updated_filename)
            
            if os.path.exists(fallback_path):
                return send_file(fallback_path, as_attachment=True, download_name=updated_filename)
            else:
                flash('❌ Updated ranking sheet not found. Please run the analysis first.')
                return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading updated ranking sheet: {str(e)}')
        return redirect(url_for('weekly_report'))


@app.route('/download_backend_file')
def download_backend_file():
    """Download the processed ROB file"""
    try:
        filename = 'ROB.xlsx'
        file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
        
        if os.path.exists(file_path):
            return send_file(file_path, as_attachment=True, download_name=filename)
        else:
            flash('❌ Processed file not found. Please process a backend file first.')
            return redirect(url_for('weekly_report'))
    except Exception as e:
        flash(f'❌ Error downloading file: {str(e)}')
        return redirect(url_for('weekly_report'))


# ============================================================================
# PROGRESSIVE GOOGLE TRENDS FUNCTIONS
# ============================================================================

# Updated Google Trends Config
GOOGLE_TRENDS_CONFIG = {
    'API_KEY': '68789844ac879558272d4e4d',  # ScrapingDog API key
    'INTEREST_THRESHOLD': 50,
    'DAYS_ABOVE_THRESHOLD': 2,
    'TERMS_TO_REMOVE': ['market', 'size', 'analysis', 'report', 'industry', 'global'],
    'REQUEST_DELAY': 2
}

class GoogleTrendsExtractor:
    def __init__(self, api_key):
        self.api_key = api_key
        self.base_url = "https://api.scrapingdog.com/google_trends"
    
    def get_values(self, keyword):
        """Get exact values for last 7 days"""
        params = {
            "api_key": self.api_key,
            "query": keyword,
            "geo": "US",           # Worldwide
            "tz": "330",         # Indian timezone (UTC+5:30)
            "date": "now 7-d",   # Last 7 days
            "data_type": "TIMESERIES"
        }
        try:
            response = requests.get(self.base_url, params=params)
            if response.status_code == 200:
                data = response.json()
                values = self.extract_values(data)
                return values
            else:
                print(f"    API Error: {response.status_code}")
                return []
        except Exception as e:
            print(f"    Network Error: {e}")
            return []
    
    def extract_values(self, data):
        """Extract values using standard timeline method"""
        values = []
        
        try:
            if 'interest_over_time' in data:
                timeline_data = data['interest_over_time'].get('timeline_data', [])
                
                for entry in timeline_data:
                    if isinstance(entry, dict) and 'values' in entry:
                        for val_item in entry['values']:
                            if isinstance(val_item, dict) and 'value' in val_item:
                                try:
                                    val = int(val_item['value'])
                                    if 0 <= val <= 100:
                                        values.append(val)
                                except (ValueError, TypeError):
                                    pass
        except Exception:
            pass
        
        return values
    
    def filter_keyword(self, keyword, threshold=2):
        """Check if keyword has threshold+ values > 50 in last 7 days"""
        values = self.get_values(keyword)
        
        if values:
            count_above_50 = sum(1 for val in values if val > 50)
            print(f"    Values: {values[:10]}{'...' if len(values) > 10 else ''} | Count >50: {count_above_50} | Threshold: {threshold}+")
            return count_above_50 >= threshold
        
        print(f"    No values retrieved | Threshold: {threshold}+")
        return False

def analyze_keywords_with_progressive_trends(keywords_data, target_count=None):
    """Analyze keywords with 2-phase Google Trends filtering: 2+ → 1+"""
    api_key = GOOGLE_TRENDS_CONFIG['API_KEY']
    extractor = GoogleTrendsExtractor(api_key)
    
    total_keywords = len(keywords_data)
    print(f"🔍 Starting 2-Phase Google Trends Analysis for {total_keywords} keywords...")
    
    if target_count:
        print(f"🎯 Target: {target_count} qualified keywords")
    else:
        print(f"🎯 Target: Not specified (will use medium 2+ criteria only)")
    
    # Additional safety check
    if total_keywords > 1500:
        print(f"⚠️ WARNING: Received {total_keywords} keywords, limiting to 1500!")
        keywords_data = keywords_data[:1500]
    
    # Initialize tracking
    qualified_keywords = []
    checked_rids = set()
    
    # Results breakdown for final report
    results_breakdown = {
        'medium_2plus': 0, 
        'relaxed_1plus': 0,
        'total_checked': 0
    }
    
    # PHASE 1: Try medium criteria first (2+ values > 50)
    print(f"\n🟡 PHASE 1: Medium Filter (2+ values > 50)")
    print(f"{'='*60}")
    
    for i, keyword_row in enumerate(keywords_data):
        try:
            original_keyword = keyword_row.get('Keywords', '')
            rid = keyword_row.get('RID', '')
            competition = keyword_row.get('Competition', '')
            search_volume = keyword_row.get('AVG. Search', 0)
            
            if not original_keyword or not rid or rid in checked_rids:
                continue
                
            clean_keyword = clean_keyword_for_trends(original_keyword)
            if not clean_keyword:
                continue
            
            print(f"[{i+1}/{len(keywords_data)}] Medium Check RID {rid}: '{clean_keyword}' [{competition}, {search_volume:,}]")
            
            is_trending = extractor.filter_keyword(clean_keyword, threshold=2)
            results_breakdown['total_checked'] += 1
            
            if is_trending:
                qualified_keywords.append({
                    'RID': rid, 
                    'keyword': original_keyword,
                    'competition': competition,
                    'search_volume': search_volume,
                    'qualified_at': '2+'
                })
                checked_rids.add(rid)
                results_breakdown['medium_2plus'] += 1
                print(f"  ✅ QUALIFIED (2+): RID {rid}")
            else:
                print(f"  ❌ Not medium enough: RID {rid}")
                
            time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                
        except Exception as e:
            print(f"  ❌ Error analyzing RID {keyword_row.get('RID', 'unknown')}: {e}")
            continue
    
    print(f"\n📊 PHASE 1 RESULTS: {results_breakdown['medium_2plus']} keywords qualified with 2+ criteria")
    
    # Check if we need to continue
    if not target_count or len(qualified_keywords) >= target_count:
        print(f"🎯 Target reached with medium criteria! Stopping here.")
        return finalize_2phase_results(qualified_keywords, results_breakdown, target_count)
    
    # PHASE 2: Try relaxed criteria (1+ values > 50) on remaining keywords
    shortage_after_medium = target_count - len(qualified_keywords)
    print(f"\n🟢 PHASE 2: Relaxed Filter (1+ values > 50)")
    print(f"Need {shortage_after_medium} more keywords...")
    print(f"{'='*60}")
    
    for i, keyword_row in enumerate(keywords_data):
        try:
            rid = keyword_row.get('RID', '')
            
            if not rid or rid in checked_rids:
                continue
                
            original_keyword = keyword_row.get('Keywords', '')
            competition = keyword_row.get('Competition', '')
            search_volume = keyword_row.get('AVG. Search', 0)
            
            clean_keyword = clean_keyword_for_trends(original_keyword)
            if not clean_keyword:
                continue
            
            print(f"[{i+1}/{len(keywords_data)}] Relaxed Check RID {rid}: '{clean_keyword}' [{competition}, {search_volume:,}]")
            
            is_trending = extractor.filter_keyword(clean_keyword, threshold=1)
            results_breakdown['total_checked'] += 1
            
            if is_trending:
                qualified_keywords.append({
                    'RID': rid, 
                    'keyword': original_keyword,
                    'competition': competition,
                    'search_volume': search_volume,
                    'qualified_at': '1+'
                })
                checked_rids.add(rid)
                results_breakdown['relaxed_1plus'] += 1
                print(f"  ✅ QUALIFIED (1+): RID {rid}")
                
                # Check if we've reached target
                if len(qualified_keywords) >= target_count:
                    print(f"🎯 Target reached! Stopping relaxed filter.")
                    break
            else:
                print(f"  ❌ Not relaxed enough: RID {rid}")
                
            time.sleep(GOOGLE_TRENDS_CONFIG['REQUEST_DELAY'])
                
        except Exception as e:
            print(f"  ❌ Error analyzing RID {keyword_row.get('RID', 'unknown')}: {e}")
            continue
    
    print(f"\n📊 PHASE 2 RESULTS: {results_breakdown['relaxed_1plus']} additional keywords qualified with 1+ criteria")
    
    return finalize_2phase_results(qualified_keywords, results_breakdown, target_count)


def finalize_2phase_results(qualified_keywords, results_breakdown, target_count):
    """Finalize and display 2-phase Google Trends results"""
    
    print(f"\n{'='*80}")
    print(f"🎯 2-PHASE GOOGLE TRENDS ANALYSIS COMPLETE")
    print(f"{'='*80}")
    
    total_qualified = len(qualified_keywords)
    
    # Display breakdown
    print(f"📊 QUALIFICATION BREAKDOWN:")
    print(f"  🟡 Medium (2+ values > 50):  {results_breakdown['medium_2plus']} keywords")
    print(f"  🟢 Relaxed (1+ values > 50): {results_breakdown['relaxed_1plus']} keywords")
    print(f"  ➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖➖")
    print(f"  📈 TOTAL QUALIFIED:          {total_qualified} keywords")
    print(f"  🔍 TOTAL CHECKED:            {results_breakdown['total_checked']} API calls")
    
    # Target analysis
    if target_count:
        if total_qualified >= target_count:
            print(f"  ✅ TARGET STATUS:            ACHIEVED ({total_qualified}/{target_count})")
        else:
            shortage = target_count - total_qualified
            print(f"  ⚠️ TARGET STATUS:            SHORT BY {shortage} ({total_qualified}/{target_count})")
            print(f"  💡 RECOMMENDATION:           Add more keywords to input file for next time")
    else:
        print(f"  ℹ️ TARGET STATUS:            No target specified")
    
    # Sort qualified keywords by priority
    if qualified_keywords:
        priority_map = {'Low': 1, 'Medium': 2, 'High': 3}
        qualified_keywords.sort(key=lambda x: (
            priority_map.get(x.get('competition', 'High'), 4),
            -x.get('search_volume', 0)
        ))
        print(f"📈 Qualified keywords sorted by priority and search volume")
    
    return qualified_keywords

def clean_keyword_for_trends(keyword):
    """Clean keyword by removing problematic terms"""
    if not keyword:
        return ""
        
    cleaned = str(keyword)
    
    # Remove terms from config
    for term in GOOGLE_TRENDS_CONFIG['TERMS_TO_REMOVE']:
        cleaned = re.sub(rf'\b{re.escape(term)}\b', '', cleaned, flags=re.IGNORECASE)
    
    # Clean up extra spaces and trim
    cleaned = ' '.join(cleaned.split()).strip()
    return cleaned
# ============================================================================
# POWER AUTOMATE ROUTES
# ============================================================================

@app.route('/wait_power_automate')
def wait_power_automate():
    """Show a waiting page for Power Automate Desktop step."""
    return render_template('wait_power_automate.html')

@app.route('/api/trigger_power_automate', methods=['POST'])
def trigger_power_automate_flow():
    """Triggers a Power Automate Desktop flow"""
    pad_exe_path = r"C:\Program Files (x86)\Power Automate Desktop\PAD.Console.Host.exe"
    flow_name = "Paid PR - Files Downloader"
    
    if not os.path.exists(pad_exe_path):
        print("Power Automate Desktop executable not found!")
        return jsonify({'status': 'error', 'message': 'PAD executable not found'})
    
    command = f'"{pad_exe_path}" -flow "{flow_name}"'
    
    try:
        result = subprocess.run(command, shell=True, check=True, text=True, capture_output=True)
        print(f"Flow triggered successfully. Output: {result.stdout}")

        time.sleep(5)
        
        flow_button_coordinates = (463, 395)
        print(f"Clicking at {flow_button_coordinates}")
        pyautogui.click(flow_button_coordinates)
        print("Flow triggered successfully.")

    except subprocess.CalledProcessError as e:
        print(f"Error triggering flow: {e.stderr}")
        return jsonify({'status': 'error', 'message': f'Flow error: {e.stderr}'})
    
    return jsonify({'status': 'success', 'message': 'Power Automate process completed.'})

# ============================================================================
#  Custom APPLICATION RUNNER
# ============================================================================
@app.route('/custom_index.html')
def custom_index():
    """Render custom index page for application"""
    return render_template('custom_index.html')



@app.route('/custom_weekly_report', methods=['GET', 'POST'])
def costum_weekly_report():
    if request.method == 'POST':
        form_type = request.form.get('form_type')
        
        if form_type == 'backend_processing':
            return handle_backend_processing()
        else:
            return handle_rid_analysis_with_domains()
    
    # GET request - show form
    return render_template('weekly_report.html', qualified_rids=None, filter_summary=None, backend_result=None)


def handle_rid_analysis_with_domains():
    """Handle RID analysis with improved domain distribution workflow"""
    try:
        print("Improved RID Analysis with Domain Distribution POST request received!")
        
        # Get filter parameters from form
        min_search_volume = int(request.form.get('min_search_volume', 5000))
        competition_level = request.form.get('competition_level', 'Low')
        analyze_trends = request.form.get('analyze_trends') == 'on'
        target_count = request.form.get('target_count')
        
        # Convert target_count to int if provided
        if target_count and target_count.strip():
            try:
                target_count = int(target_count)
                if target_count <= 0:
                    raise ValueError("Target count must be positive")
            except ValueError:
                flash('❌ Please enter a valid target word count!')
                return redirect(request.url)
        else:
            target_count = None
        
        print(f"User Filters: Search >= {min_search_volume}, Competition = {competition_level}")
        print(f"Google Trends: {'Enabled' if analyze_trends else 'Disabled'}")
        print(f"Target Count: {target_count if target_count else 'Not specified'}")
        
        # Validate form inputs
        if not min_search_volume or min_search_volume < 0:
            flash('❌ Please enter a valid minimum search volume!')
            return redirect(request.url)
            
        if not competition_level:
            flash('❌ Please select a competition level!')
            return redirect(request.url)
        
        # Handle file uploads
        ranking_file = request.files.get('ranking_file')
        if not ranking_file or ranking_file.filename == '':
            flash('❌ Please select a ranking Excel file!')
            return redirect(request.url)

        if not allowed_file(ranking_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ranking sheet!')
            return redirect(request.url)

        rob_file = request.files.get('cleaned_rob_file')
        if not rob_file or rob_file.filename == '':
            flash('❌ Please select a cleaned ROB Excel file!')
            return redirect(request.url)

        if not allowed_file(rob_file.filename):
            flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed for ROB file!')
            return redirect(request.url)

        # Save uploaded files
        ranking_filename = secure_filename(ranking_file.filename)
        ranking_path = os.path.join(app.config['UPLOAD_FOLDER'], ranking_filename)
        ranking_file.save(ranking_path)
        
        rob_filename = secure_filename(rob_file.filename)
        rob_path = os.path.join(app.config['UPLOAD_FOLDER'], rob_filename)
        rob_file.save(rob_path)
        
        # Process files with improved workflow
        result_summary = process_dual_files_improved_workflow(
            ranking_path, rob_path, min_search_volume, competition_level, analyze_trends, target_count
        )
        
        # Display comprehensive results
        if result_summary['success']:
            # Main success message
            flash(f'✅ Success! Found {result_summary["qualified_rids_count"]} qualified RIDs')
            
            # Show original counts
            original_stats = result_summary.get("original_stats", {})
            if original_stats:
                flash(f'📊 Original Data: Total={original_stats.get("total", 0)}, HC={original_stats.get("HC", 0)}, ICT={original_stats.get("ICT", 0)}, CMFE={original_stats.get("CMFE", 0)}')
            
            # Show target vs achieved
            target_info = result_summary.get("target_info", {})
            if target_info:
                flash(f'🎯 Target: {target_info.get("target_total", 0)} words ({target_info.get("target_per_domain", 0)} per domain)')
                flash(f'🎯 Achieved: HC={target_info.get("achieved_HC", 0)}, ICT={target_info.get("achieved_ICT", 0)}, CMFE={target_info.get("achieved_CMFE", 0)}')
            
            # Show shortage information
            shortage_info = result_summary.get("shortage_info", {})
            if shortage_info.get("has_shortage"):
                for domain, shortage in shortage_info.get("shortages", {}).items():
                    if shortage > 0:
                        flash(f'⚠️ {domain} Domain Shortage: {shortage} words (filled from other domains)')
                
                total_shortage = shortage_info.get("total_shortage", 0)
                if total_shortage > 0:
                    flash(f'💡 Recommendation: Add at least {total_shortage} more {", ".join(shortage_info.get("short_domains", []))} domain words next time')
            
            # Show search volume adjustments
            search_adjustments = result_summary.get("filter_summary", {}).get("search_adjustments", {})
            if search_adjustments.get("adjusted"):
                original_sv = search_adjustments.get("original_volume", min_search_volume)
                final_sv = search_adjustments.get("final_volume", min_search_volume)
                attempts = search_adjustments.get("attempts", [])
                
                if len(attempts) > 1:
                    if final_sv == 0:
                        flash(f'🔄 Auto-adjusted search volume: {original_sv:,} → No Filter (tried: {", ".join([f"{v:,}" if v > 0 else "No Filter" for v in attempts])})')
                    else:
                        flash(f'🔄 Auto-adjusted search volume: {original_sv:,} → {final_sv:,} (tried: {", ".join([f"{v:,}" if v > 0 else "No Filter" for v in attempts])})')
                else:
                    if final_sv == 0:
                        flash(f'📈 Used search volume: No Filter (all keywords)')
                    else:
                        flash(f'📈 Used search volume: {final_sv:,}')
            else:
                flash(f'📈 Used search volume: {min_search_volume:,}')
           
           # Show Google Trends progressive results
            trends_info = result_summary.get("trends_info", {})
            if trends_info.get("enabled") and trends_info.get("breakdown"):
                breakdown = trends_info["breakdown"] 
                flash(f'🔍 Google Trends: {breakdown["medium_2plus"]} (2+), {breakdown["relaxed_1plus"]} (1+) = {breakdown["total_qualified"]} total')
            # Show final results vs target
            if target_count:
                if result_summary["qualified_rids_count"] < target_count:
                    final_shortage = target_count - result_summary["qualified_rids_count"]
                    flash(f'⚠️ Only {result_summary["qualified_rids_count"]} words qualifying out of {target_count} requested')
                    flash(f'💡 Please add {final_shortage} more words to your input file for next time')
                elif result_summary["qualified_rids_count"] == target_count:
                    flash(f'🎯 Perfect! Achieved exactly your target of {target_count} words')
                else:
                    flash(f'✨ Exceeded target! Got {result_summary["qualified_rids_count"]} out of {target_count} requested')
            
            # ROB extraction results
            flash(f'✅ Extracted {result_summary["matched_rob_rows"]} matching ROB rows')
            flash(f'📁 Weekly ROB.xlsx saved to Desktop/RPA folder!')
            
            # Updated ranking sheet info
            if result_summary.get("updated_ranking_path"):
                flash(f'📄 Updated ranking sheet saved and ready for download!')
            
        else:
            flash(f'❌ Error: {result_summary.get("error", "Unknown error")}')
            result_summary = None
        
        # Clean up uploaded files
        try:
            os.remove(ranking_path)
            os.remove(rob_path)
        except Exception as cleanup_error:
            print(f"Warning: Could not clean up files: {cleanup_error}")
        
        # Render template with results
        return render_template('weekly_report.html', 
                              qualified_rids=result_summary.get('qualified_rids', []) if result_summary else [],
                              filter_summary=result_summary.get('filter_summary', {}) if result_summary else {},
                              backend_result=None,
                              rob_extraction_result=result_summary)
        
    except Exception as e:
        print(f"Error: {e}")
        flash(f'❌ Error processing files: {str(e)}')
        return redirect(request.url)
 

@app.route('/custom_choice')
def custom_choice():
    """Custom choice page for CMI/WMR selection"""
    return render_template('custom_choice.html')



@app.route('/custom_cmi_cta', methods=['GET', 'POST'])
def custom_cmi_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for CMI CTA generation
            result = process_cmi_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 CMI automation started!')
                
                # Start CMI automation in background
                threading.Thread(target=run_cmi_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_cmi_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_cmi_cta.html')

def process_cmi_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for CMI CTA generation"""
    try:
        # Read the file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows

        # Step 1: Extract top N rows for CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_cmi_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated Custom_weekly_ROB.xlsx - Removed {extract_count} extracted keywords")
            print(f"✅ Custom_weekly_ROB.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: Custom_weekly_ROB.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"Error in process_cmi_cta_file: {e}")
        return {
            'success': False,
            'error': str(e)
        }


def run_cmi_automation():
    """Run CMI automation using Selenium"""
    try:
        print("Starting CMI automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time
        
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")
        
        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.coherentmarketinsights.com/cmisitmanup/index.php')
        
        username_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
        )
        username_input.send_keys('Auto_Ops_Team')
        
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
        )
        password_input.send_keys('kDp7%8^03Ib')
        
        signup_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
        )
        signup_click.click()
        
        custom_insights_click = WebDriverWait(driver, 10).until(
            EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
        )
        custom_insights_click.click()
        
        print("CMI automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"CMI automation error: {e}")


@app.route('/custom_wmr_cta', methods=['GET', 'POST'])
def custom_wmr_cta():
    if request.method == 'POST':
        try:
            # Get form data (only extract_count now since user info fields are removed)
            extract_count = int(request.form.get('extract_count', 200))

            # Validate required fields (only extract_count and file now)
            if not extract_count or extract_count < 1:
                flash('❌ Please enter a valid number of keywords to extract!')
                return redirect(request.url)

            file = request.files.get('file')
            if not file or file.filename == '':
                flash('❌ Keywords file is required!')
                return redirect(request.url)

            if not allowed_file(file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) and CSV files are allowed!')
                return redirect(request.url)

            # Use secure_filename to avoid path issues
            filename = secure_filename(file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            file.save(input_path)

            # Process the file for WMR CTA generation
            result = process_wmr_cta_file(input_path, extract_count)
            
            if result['success']:
                flash(f'✅ Successfully processed {result["extracted_count"]} keywords!')
                flash(f'📁 File saved: {result["filename"]}')
                flash('🤖 WMR automation started!')
                
                # Start WMR automation in background
                threading.Thread(target=run_wmr_automation).start()
            else:
                flash(f'❌ Error: {result["error"]}')

            return render_template('custom_wmr_cta.html')

        except ValueError as ve:
            flash('❌ Please enter a valid number for keywords to extract!')
            return redirect(request.url)
        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_wmr_cta.html')


def process_wmr_cta_file(file_path, extract_count):
    """Process Custom Weekly ROB file for WMR CTA generation"""
    try:
        # Read the file
        if file_path.endswith('.csv'):
            df_original = pd.read_csv(file_path)
        else:
            df_original = pd.read_excel(file_path, engine='openpyxl')

        total_rows = len(df_original)
        
        if total_rows < extract_count:
            extract_count = total_rows

        # Step 1: Extract top N rows for WMR CTA generation
        extracted_rows = df_original.head(extract_count).copy()
        
        # Step 2: Get remaining rows (original minus extracted)
        remaining_rows = df_original.iloc[extract_count:].copy()
        
        # Create filename with current date for extracted data
        today = datetime.today()
        extracted_filename = f"custom_reports_wmr_{today.year}_{today.month:02d}_{today.day:02d}.xlsx"
        
        # Save extracted data to Desktop/RPA folder
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        # Save extracted keywords file
        extracted_output_path = os.path.join(rpa_folder, extracted_filename)
        extracted_rows.to_excel(extracted_output_path, index=False)
        
        # Step 3: Update the original Custom_weekly_ROB.xlsx file with remaining data
        custom_weekly_rob_path = os.path.join(rpa_folder, 'Custom_weekly_ROB.xlsx')
        
        if os.path.exists(custom_weekly_rob_path):
            # Update the original file with remaining rows
            remaining_rows.to_excel(custom_weekly_rob_path, index=False)
            print(f"✅ Updated Custom_weekly_ROB.xlsx - Removed {extract_count} extracted keywords for WMR")
            print(f"✅ Custom_weekly_ROB.xlsx now contains {len(remaining_rows)} remaining keywords")
        else:
            print(f"⚠️ Warning: Custom_weekly_ROB.xlsx not found at {custom_weekly_rob_path}")
        
        return {
            'success': True,
            'extracted_count': extract_count,
            'remaining_count': len(remaining_rows),
            'total_count': total_rows,
            'filename': extracted_filename,
            'output_path': extracted_output_path,
            'updated_original': os.path.exists(custom_weekly_rob_path)
        }
        
    except Exception as e:
        print(f"Error in process_wmr_cta_file: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def run_wmr_automation():
    """Run WMR automation using Selenium with your provided code"""
    try:
        print("Starting WMR automation...")
        
        from webdriver_manager.chrome import ChromeDriverManager
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        import time

        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        #options.add_argument("--headless")  # Run headlessly for deployment
        options.add_argument("--start-maximized")

        # Initialize WebDriver
        cService = Service(executable_path=chromedriver_path)
        driver = webdriver.Chrome(service=cService, options=options)
        driver.get('https://www.worldwidemarketreports.com/imanagereports')
                
        username_input = WebDriverWait(driver, 10).until(
             EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[1]/input'))
            )
        username_input.send_keys('Auto_Ops_Team')
                
        password_input = WebDriverWait(driver, 10).until(
            EC.presence_of_element_located((By.XPATH,'//*[@id="adlogin"]/div[2]/input'))
            )
        password_input.send_keys('M9b@0j9Y28O')
                
        login_click = WebDriverWait(driver, 10).until(
          EC.element_to_be_clickable((By.XPATH,'//*[@id="adlogin"]/div[3]/div/button'))
            )
        login_click.click()
                
        custom_insights_click = WebDriverWait(driver, 10).until(
           EC.element_to_be_clickable((By.XPATH,'/html/body/div/aside/section/ul/li[3]/a/span[1]'))
            )
        custom_insights_click.click()
        
        print("WMR automation completed successfully!")
        
        # Keep browser open for now - you can modify this behavior
        time.sleep(10)
        # driver.quit()  # Uncomment to close browser automatically
        
    except Exception as e:
        print(f"WMR automation error: {e}")


@app.route('/custom_content_generation_choice')
def custom_content_generation_choice():
    """Custom content generation choice page"""
    return render_template('custom_content_generation_choice.html')

# Add these imports at the top if not already present
import openai
from docx import Document
import re

# Configure OpenAI (add your API key)
OPENAI_API_KEY = "sk-proj-Vl99-wfm2By_9pPYUKtnx0kXubr3cAdQ18h4TpM3H4WA_q
@app.route('/custom_ai_content', methods=['GET', 'POST'])
def custom_ai_content():
    if request.method == 'POST':
        try:
            # Handle file upload
            cta_file = request.files.get('cta_file')
            if not cta_file or cta_file.filename == '':
                flash('❌ CTA excel file is required!')
                return redirect(request.url)

            if not allowed_file(cta_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(cta_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            cta_file.save(input_path)

            # Process the file for AI content generation
            result = process_ai_content_generation(input_path)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["articles_created"]} AI articles!')
                flash(f'📁 Articles saved to Desktop/RPA folder')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_ai_content.html')

        except Exception as e:
            flash(f'❌ Error processing file: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_ai_content.html')

def clean_title(title):
    """Remove 'Market' and related words from title"""
    # Remove common market-related terms
    market_terms = [
        r'\bmarket\b', r'\bMarket\b', r'\bMARKET\b',
        r'\bmarket size\b', r'\bMarket Size\b',
        r'\bmarket analysis\b', r'\bMarket Analysis\b',
        r'\bmarket research\b', r'\bMarket Research\b',
        r'\bmarket report\b', r'\bMarket Report\b',
        r'\bmarket study\b', r'\bMarket Study\b'
    ]
    
    cleaned_title = title
    for term in market_terms:
        cleaned_title = re.sub(term, '', cleaned_title, flags=re.IGNORECASE)
    
    # Clean up extra spaces and punctuation
    cleaned_title = re.sub(r'\s+', ' ', cleaned_title).strip()
    cleaned_title = re.sub(r'^[-\s]+|[-\s]+$', '', cleaned_title)
    
    return cleaned_title

def generate_article_with_openai(clean_title, promo_link, sample_link):
    """Generate article using OpenAI API"""
    try:
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        
        prompt = f"""
We are coherent market insights and going to publish this data on third part PR website, thus do not mention any competitors name (market intelligence firm, market research, consulting and data and analytics companies), and take ownership for market share, size, growth rate as this research done by our own research team however any real instances and stats you can use reference such as company or organisation or others as mentioned in detail explanation. Write the blog covering the below pointers: Start content with market size and overview paragraph (do not add heading and do not use bullet for this section), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here, ➤Market Taxonomy and Regional Coverage of Report, ➤Leading Companies of the Market, ➔ add Second CTA link here, ➤Key Growth Drivers Fueling Market Expansion, ➤Key Reasons for Buying the (insert market name here) Report ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ❓ Frequently Asked Questions, All these pointers should act as a heading to respective paragraphs, do not miss any bullet for above given. Consider the following points while generating content: Flow of information, all given secondary keywords must be covered, and there should be a heading for each paragraph or bullet pointers. I need this content to be very presentable format thus make sure there should be space after and before of each heading or CTA links (Strategic Actionable Insights for the Market, first CTA link, Market Taxonomy and Regional coverage of Report, Leading Companies of the Market, Second CTA link, Key Growth Drivers Fueling Market Expansion, Key Reasons for Buying the (insert market name here) Report, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, and Frequently Asked Questions), this will increase the readability. Cover content in in bullet pointers whenever possible each paragraph should be short. Our readers are already experts in the field, so always try to generate content that provides unique insights and value addition for experts. Thus, while generating each piece of content, it should be data-backed with actual instances and stats from recent years 2025 and 2024, covering maximum stats that increase the authenticity and are enough to support or make decisions based upon reading this blog. Do not add generic content that is already known to readers, such as definitions, advantages, disadvantages, or other generic content. Flow of information should be as below. Start the blog with writing 5 to 7 sentence ling paragraph start content with talking about Industry in 2 to 3 sentences and should be matches with heading of the blog. followed by 2 to 4 sentence on market size and CAGR formatted as The Global (Insert Market Name) Market size is estimated to be valued at USD (Identify and Insert Market Value for 2025) in 2025 (market size) and is expected to reach USD (Identify and Insert Market Value for 2032) by 2032 (market forecast), exhibiting a compound annual growth rate (CAGR) of (Identify and Insert CAGR)% from 2025 to 2032. Do not add heading and do not use bullet for this section. (Make sure all values such as market size, CAGR, key players should be carefully identified with research approach). Then Strategic Actionable Insights for the Market: In Strategic Actionable Insights for the Market, cover 3 to 4 Strategic Actionable Insights for the Market in 10 to 12 sentences each insights comprises of 2 t 3 sentences covering supporting data for each insights. Insights could be essential quantitative indicators that lead to market size. For example, supply-side indicators (ex. Production Capacity, Pricing, and Exports. but not necessary to add if not matches to given market), demand-side indicators (ex. Pricing, Imports, Various Use Cases across industries. but not necessary to add if not matches to given market), micro-indicators, nano-size indicators—focus more on the quantitative aspects. Each Strategic Actionable Insights for the Market must have two sentence stats or actual instance examples from the recent year to support each point given in Strategic Actionable Insights for the Market, so that each given point look complete and meaningful. Then First CTA link. Next part is Market Taxonomy and Regional coverage of Report where enlist the all subsegment under each segment categories and fragment region into given format. Identify Comprehensive Market Taxonomy of the Report: » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. » By Segment 1: subsegment 1, subsegment 2, subsegment 3, subsegment 4, and Others. Regional and Country Analysis: » North America: U.S. and Canada » Latin America: Brazil, Argentina, Mexico, and Rest of Latin America » Europe: Germany, U.K., Spain, France, Italy, Benelux, Denmark, Norway, Sweden, Russia, and Rest of Europe » Asia Pacific: China, Taiwan, India, Japan, South Korea, Indonesia, Malaysia, Philippines, Singapore, Australia, and Rest of Asia Pacific » Middle East & Africa: Bahrain, Kuwait, Oman, Qatar, Saudi Arabia, United Arab Emirates, Israel, South Africa, North Africa, Central Africa, and Rest of MEA. Then Leading Companies of the Market: Identify and Enlist 12 to 20 highly relevant Leading Companies of the Market for the given market. Furthermore, add 2-3 statements on competitive strategies adopted by a few Leading Companies of the Market, mentioning actual strategies and entities involved along with the actual outcome. Then Add Second CTA link. Key Reasons for Buying the (insert market name here) Report, and its exact content as shared in data. Key Growth Drivers Fueling Market Expansion: Growth factor heading and short paragraph (3-4 Key Growth Drivers Fueling Market Expansion covered under 10 to 12 sentences) with supporting stats or examples from the recent year in the content, each factors should be covered in two to three sentences thus entire Key Growth Drivers Fueling Market Expansion content will be covered in 10 to 12 sentences long. No sub bullet is needed in Growth Factor. Then Emerging Trends and Market Shift: Market Trend heading and short paragraphs with supporting stats or examples from the recent year in the content (No bullet needed for as opportunity are written in paragraph format). Then High-Impact Market Opportunities by Segment and Region: Provide 3 to 4 High-Impact Market Opportunities by Segment and Region, 2-3 opportunities based upon segment and one opportunity based upon region in a paragraph format. Each opportunity described in two to three sentences and supported by actual data. For each opportunity, identify a single segment and provide a short description of the opportunity within that segment. Similarly, highlight one region, along with a brief description of the opportunity in each regional market (No bullet needed for as opportunity are written in paragraph format). Make sure to not to mention our organization name or relevant terms anywhere in the output such as coherent market insights or our analyst team or our research team Given Market Name and Data:
Input of different market
 ➤Key Reasons for Buying the (insert market name here) Report: ✦ Comprehensive analysis of the changing competitive landscape ✦ Assists in decision-making processes for the businesses along with detailed strategic planning methodologies ✦ The report offers forecast data and an assessment of the (insert market name here) ✦ Helps in understanding the key product segments and their estimated growth rate ✦ In-depth analysis of market drivers, restraints, trends, and opportunities ✦ Comprehensive regional analysis of the (insert market name here) ✦ Extensive profiling of the key stakeholders of the business sphere ✦ Detailed analysis of the factors influencing the growth of the (insert market name here). From an SEO perspective, we need to cover all given keywords from the list below. However, they should appear naturally so that the content flow looks natural for the reader. Keyword List: market share, market size, market research, market insights, Emerging Trends and Market Shift, High-Impact Market Opportunities by Segment and Region, market challenges, market growth, market forecast, market companies, market players, market analysis, market drivers, market restraints, market scope, market dynamics, market segments, market report, market growth strategies, market revenue, industry size, industry share, industry trends, and business growth, furthermore - Market size and market report, market revenue, market share, trends keywords are mandatory to be added twice in content. In addition to the above requirement, in 5 places, add the actual market name along with the above keywords so that long-tail keywords will be generated. These long-tail keywords are market name + size, market name + report, market name + revenue, market name + share, market name + trends. Make sure all given keywords are naturally fit, do not try to infuse forcefully, flow of information should be natural and meaningful, furthermore make sure spelling and structure of sentences from generated output are grammatically correct. Furthermore, based on the market name, create a set of Frequently Asked Questions that are highly relevant and customized to the specific market. The sample Frequently Asked Questions below are for understanding purposes only. For the given market, questions can be completely replaced. However, please tailor the actual questions to the market name and the insights provided in the report: 1. Who are the dominant players in the (Market Name) market? 2. What will be the size of the (Market Name) market in the coming years? 3. Which end users industry has the largest growth opportunity? 4. How will market development trends evolve over the next five years? 5. What is the nature of the competitive landscape and challenges in the (Market Name) market? 6. What go-to-market strategies are commonly adopted in the (Market Name) market? Make sure to answer to all FAQs. In the case of country-level markets, please exclude the word 'Global' and High-Impact Market Opportunities by Segment and Region where other regions are mentioned. Make sure to add catchy bullet in generated output. I have shared the reference bullet with you. Make sure to add this bullet. For heading use these bullet- ), ➤Strategic Actionable Insights for the Market, ➔ add first CTA link here ➤Market Taxonomy and Regional coverage of Report, ➤Leading Companies of the Market, ➔ Inserted Second CTA link, ➤Key Reasons for Buying the (insert market name here) Report, ➤Key Growth Drivers Fueling Market Expansion, ➤ Emerging Trends and Market Shift, ➤High-Impact Market Opportunities by Segment and Region, and ❓ Frequently Asked Questions. Make sure do not miss any bullet including CTA bullet which is ➔. For subpointers under main headings use bullets which is in reference as provided- Strategic Actionable Insights for the Market ●, Market Taxonomy and Regional coverage of Report●, Leading Companies of the Market●. Make sure to use these bullets for given subpointers. Ensure proper bullet formatting so that each point is marked with a single bullet only, and avoid placing two bullets adjacent to each other.
"""
        
        response = client.chat.completions.create(
            model="gpt-4.1-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research writer specializing in industry analysis articles."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7
        )
        
        return response.choices[0].message.content
        
    except Exception as e:
        print(f"OpenAI error: {e}")
        return "Error generating article content. Please try again later."


def save_article_as_doc(article_content, clean_title):
    """Save article as .doc file"""
    try:
        # Create Word document
        doc = Document()
        
        # Add title
        title_paragraph = doc.add_heading(f"{clean_title} - Market Analysis", level=1)
        
        # Add content paragraphs
        paragraphs = article_content.split('\n\n')
        for paragraph in paragraphs:
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        # Create filename
        today = datetime.today()
        safe_title = re.sub(r'[^\w\s-]', '', clean_title.lower())
        safe_title = re.sub(r'[-\s]+', '_', safe_title)
        filename = f"{safe_title}_cmi_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save to Desktop/RPA
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        rpa_folder = os.path.join(desktop_path, "RPA")
        
        if not os.path.exists(rpa_folder):
            os.makedirs(rpa_folder)
        
        file_path = os.path.join(rpa_folder, filename)
        doc.save(file_path)
        
        return True, filename
        
    except Exception as e:
        print(f"Error saving document: {e}")
        return False, str(e)

def process_ai_content_generation(file_path):
    """Process CTA excel file and generate AI articles"""
    try:
        print(f"\n=== PROCESSING AI CONTENT GENERATION ===")
        print(f"File: {file_path}")
        
        # Read the excel file
        df = pd.read_excel(file_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        articles_created = 0
        
        # Process each row
        for index, row in df.iterrows():
            try:
                original_title = str(row['KEYWORD'])
                promo_link = str(row['PROMOBUY'])
                sample_link = str(row['SAMPLECOPY'])
                
                print(f"\n[{index+1}/{len(df)}] Processing: {original_title}")
                
                # Clean the title - FIXED: Use different variable name
                cleaned_title = clean_title(original_title)  # ✅ FIXED!
                print(f"Cleaned title: {cleaned_title}")
                
                # Generate article using OpenAI
                print("Generating article with OpenAI...")
                article_content = generate_article_with_openai(cleaned_title, promo_link, sample_link)
                
                # Save as .doc file
                success, filename = save_article_as_doc(article_content, cleaned_title)
                
                if success:
                    print(f"✅ Article saved: {filename}")
                    articles_created += 1
                else:
                    print(f"❌ Failed to save article: {filename}")
                
                # Small delay to avoid API rate limits
                time.sleep(1)
                
            except Exception as e:
                print(f"❌ Error processing row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'articles_created': articles_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in AI content generation: {e}")
        return {
            'success': False,
            'error': str(e)
        }


@app.route('/custom_template_content', methods=['GET', 'POST'])
def custom_template_content():
    if request.method == 'POST':
        # Handle template-based content generation
        # Similar to custom_ai_content but using predefined templates
        pass
    return render_template('custom_template_content.html')

@app.route('/custom_wmr_templates', methods=['GET', 'POST'])
def custom_wmr_templates():
    """WMR specific template selection and processing"""
    if request.method == 'POST':
        # Handle WMR template processing
        return handle_wmr_template_processing()
    return render_template('custom_wmr_templates.html')

@app.route('/custom_cmi_templates', methods=['GET', 'POST'])  
def custom_cmi_templates():
    """CMI specific template selection and processing"""
    if request.method == 'POST':
        # Handle CMI template processing
        return handle_cmi_template_processing()
    return render_template('custom_cmi_templates.html')


#NEW #NEW 
@app.route('/custom_cmi_template_processor', methods=['GET', 'POST'])
def custom_cmi_template_processor():
    if request.method == 'POST':
        try:
            # Get the uploaded file and selected template info
            template_file = request.files.get('template_file')
            selected_category = request.form.get('selected_category', '')
            selected_template = request.form.get('selected_template', '')
            
            if not template_file or template_file.filename == '':
                flash('❌ Please upload an Excel file!')
                return redirect(request.url)

            if not allowed_file(template_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(template_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            template_file.save(input_path)

            # Process the templates
            result = process_cmi_templates_with_excel(input_path, selected_category, selected_template)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["files_created"]} template files!')
                flash(f'📁 Files saved in respective domain folders')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_cmi_templates.html')

        except Exception as e:
            flash(f'❌ Error processing templates: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_cmi_templates.html')

def process_cmi_templates_with_excel(excel_path, selected_category='', selected_template=''):
    """Process Excel file and generate templates for each row"""
    try:
        print(f"\n=== PROCESSING CMI TEMPLATES ===")
        print(f"Excel file: {excel_path}")
        
        # Read the Excel file
        df = pd.read_excel(excel_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY', 'Category']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        files_created = 0
        base_rpa_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        
        # Process each row
        for index, row in df.iterrows():
            try:
                keyword = str(row['KEYWORD']).strip()
                category = str(row['Category']).strip()
                
                
                print(f"\n[{index+1}/{len(df)}] Processing: {keyword} (Category: {category})")
                
                # Validate category
                if category not in ['HC', 'CMFE', 'ICT']:
                    print(f"❌ Invalid category: {category}. Skipping.")
                    continue
                
                # Get random template from domain folder
                template_path = get_random_template_from_domain(base_rpa_path, category)
                if not template_path:
                    print(f"❌ No templates found in {category} folder")
                    continue
                
                print(f"Selected template: {os.path.basename(template_path)}")
                
                # Process the template
                success, output_path = process_single_template(template_path, row, category, base_rpa_path)
                
                if success:
                    print(f"✅ Template generated: {os.path.basename(output_path)}")
                    files_created += 1
                else:
                    print(f"❌ Failed to generate template for: {keyword}")
                
            except Exception as e:
                print(f"❌ Error processing row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'files_created': files_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in CMI template processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_random_template_from_domain(base_path, domain):
    """Get a random .doc template from the specified domain folder"""
    try:
        domain_path = os.path.join(base_path, domain)
        
        if not os.path.exists(domain_path):
            print(f"❌ Domain folder not found: {domain_path}")
            return None
        
        # Find all .doc files in the domain folder
        doc_files = [f for f in os.listdir(domain_path) if f.lower().endswith('.doc') or f.lower().endswith('.docx')]
        
        if not doc_files:
            print(f"❌ No .doc files found in {domain_path}")
            return None
        
        # Select random template
        import random
        selected_template = random.choice(doc_files)
        template_path = os.path.join(domain_path, selected_template)
        
        print(f"Random template selected: {selected_template}")
        return template_path
        
    except Exception as e:
        print(f"Error getting random template: {e}")
        return None
    
def replace_placeholders_in_docx(excel_row, doc):
    # Define the placeholders and their corresponding Excel columns
    placeholder_mapping = {
        'KEYWORD': 'KEYWORD', 
        'PROMOBUY': 'PROMOBUY',
        'SAMPLECOPY': 'SAMPLECOPY'
    }

    # Iterate over paragraphs in the Word document
    for para in doc.paragraphs:
        for placeholder, column_name in placeholder_mapping.items():
            if placeholder in para.text:
                # Get the value from the Excel row for the placeholder
                value = str(excel_row[column_name]).strip()
                # Check if it's the 'KEYWORD' placeholder and remove the word "Market"
                if placeholder == 'KEYWORD' and "Market" in value:
                    value = value.replace("Market", "").strip()  # Remove "Market" from KEYWORD
                
                # Replace the placeholder with the actual value
                para.text = para.text.replace(placeholder, value)

    # Iterate over tables in the Word document
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for placeholder, column_name in placeholder_mapping.items():
                    if placeholder in cell.text:
                        # Get the value from the Excel row for the placeholder
                        value = str(excel_row[column_name]).strip()
                        # Check if it's the 'KEYWORD' placeholder and remove the word "Market"
                        if placeholder == 'KEYWORD' and "Market" in value:
                            value = value.replace("Market", "").strip()  # Remove "Market" from KEYWORD
                        
                        # Replace the placeholder with the actual value
                        cell.text = cell.text.replace(placeholder, value)


def process_single_template(template_path, row_data, category, base_rpa_path):
    try:
        today = datetime.today()
        
        # Generate output filename using the 'KEYWORD' and category
        keyword = str(row_data['KEYWORD']).strip()
        # Remove "Market" from the KEYWORD value
        keyword = keyword.replace("Market", "").strip()  # Remove "Market" from KEYWORD
        safe_keyword = "".join(c for c in keyword if c.isalnum() or c in (' ', '_')).strip()
        
        filename = f"{category}_{safe_keyword}_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save generated files in a separate "Generated" subfolder
        domain_folder = os.path.join(base_rpa_path, category, "Generated")
        if not os.path.exists(domain_folder):
            os.makedirs(domain_folder)
        
        output_path = os.path.join(domain_folder, filename)
        
        # Process the template
        success = process_template_with_formatting(template_path, row_data, output_path, keyword)
        
        if success:
            return True, output_path
        else:
            return False, None
            
    except Exception as e:
        print(f"Error processing single template: {e}")
        return False, None

def process_template_with_formatting(template_path, row_data, output_path, keyword):
    try:
        # Open Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open the template document
        doc = word.Documents.Open(template_path)
        
        # Prepare replacement data
        today = datetime.today()
        replacements = get_replacement_data(row_data, today)
        
        print(f"Applying {len(replacements)} replacements to template")
        
        # Replace placeholders while preserving formatting
        for placeholder, replacement in replacements.items():
            if replacement:  # Only replace if we have a value
                # Check if it's 'KEYWORD' and remove "Market"
                if placeholder == 'KEYWORD':
                    replacement = replacement.replace("Market", "").strip()  # Remove "Market"
                
                replace_text_preserve_formatting(doc, placeholder, str(replacement))
        
        # STEP 2: Replace OpenAI placeholders (ADD THIS SECTION)
        category = str(row_data.get('Category', '')).strip()
        replace_openai_placeholders(doc, keyword, category)

        # Save the document with preserved formatting
        doc.SaveAs2(output_path, FileFormat=0)  # 0 = Word 97-2003 format
        doc.Close()
        word.Quit()
        
        return True
        
    except Exception as e:
        print(f"Error processing template with formatting: {e}")
        try:
            word.Quit()
        except:
            pass
        return False


def get_replacement_data(row_data, today):
    """Get replacement data from Excel - simplified for 3 placeholders only"""
    
    # DEBUG: Print what we're working with
    print(f"DEBUG - Excel row keys: {list(row_data.keys())}")
    print(f"DEBUG - Raw values:")
    print(f"  KEYWORD: '{row_data.get('KEYWORD')}'")
    print(f"  PROMOBUY: '{row_data.get('PROMOBUY')}'") 
    print(f"  SAMPLECOPY: '{row_data.get('SAMPLECOPY')}'")
    
    replacements = {
        'KEYWORD': str(row_data.get('KEYWORD', '')),
        'PROMOBUY': str(row_data.get('PROMOBUY', '')), 
        'SAMPLECOPY': str(row_data.get('SAMPLECOPY', ''))
    }
    
    print(f"DEBUG - Final replacements: {replacements}")
    return replacements

def read_template_content(template_path):
    """Read content from template file"""
    try:
        if template_path.lower().endswith('.docx'):
            # Read .docx file
            doc = Document(template_path)
            content = []
            for paragraph in doc.paragraphs:
                content.append(paragraph.text)
            return '\n'.join(content)
        
        elif template_path.lower().endswith('.doc'):
            # Read .doc file using win32com
            word = win32com.client.Dispatch("Word.Application")
            word.Visible = False
            doc = word.Documents.Open(template_path)
            content = doc.Content.Text
            doc.Close()
            word.Quit()
            return content
        
        else:
            # Try reading as text file
            with open(template_path, 'r', encoding='utf-8') as f:
                return f.read()
                
    except Exception as e:
        print(f"Error reading template {template_path}: {e}")
        return None

def replace_text_preserve_formatting(doc, find_text, replace_text):
    try:
        print(f"DEBUG - Searching for: '{find_text}' to replace with: '{replace_text}'")
        
        replacements_made = 0
        
        # Use Selection object for better formatting preservation
        selection = doc.Application.Selection
        
        # Start from beginning of document
        selection.HomeKey(Unit=6)  # wdStory - go to start of document
        
        # Loop to find and replace all instances
        while True:
            # Use Find to locate the text
            found = selection.Find.Execute(
                FindText=find_text,
                Forward=True,
                Wrap=0,  # wdFindStop - don't wrap around
                MatchCase=False,
                MatchWholeWord=False
            )
            
            if found:
                print(f"DEBUG - Found '{find_text}' at position {selection.Start}")
                # Replace the selected text (preserves formatting of surrounding text)
                selection.TypeText(replace_text)
                replacements_made += 1
                print(f"DEBUG - Replaced occurrence #{replacements_made}")
            else:
                break  # No more instances found
        
        print(f"DEBUG - Total replacements made: {replacements_made}")
        
        # Verify replacement worked
        final_content = doc.Content.Text
        if find_text in final_content:
            print(f"DEBUG - ❌ '{find_text}' still exists after replacement!")
        else:
            print(f"DEBUG - ✅ '{find_text}' successfully removed!")
            
        if replace_text in final_content:
            print(f"DEBUG - ✅ New text '{replace_text}' found in document!")
        else:
            print(f"DEBUG - ❌ New text '{replace_text}' NOT found!")
        
        print(f"Formatting-preserving replacement completed for '{find_text}'")

    except Exception as e:
        print(f"Error in formatting-preserving replacement '{find_text}': {e}")

def clean_content_for_xml(content):
    """Clean content to make it XML compatible"""
    import re
    
    if not content:
        return ""
    
    # Convert to string if not already
    content = str(content)
    
    # Remove NULL bytes and control characters (except \t, \n, \r)
    content = re.sub(r'[\x00-\x08\x0B\x0C\x0E-\x1F\x7F]', '', content)
    
    # Remove or replace other problematic characters
    # Remove excessive whitespace
    content = re.sub(r'\n\s*\n\s*\n+', '\n\n', content)  # Multiple empty lines to double
    content = re.sub(r'[ \t]+', ' ', content)  # Multiple spaces to single space
    
    # Remove any remaining XML-incompatible characters
    # Keep only printable ASCII + extended characters, newlines, tabs
    content = ''.join(char for char in content if ord(char) >= 32 or char in '\n\r\t')
    
    # Ensure proper encoding
    try:
        # Encode and decode to handle any encoding issues
        content = content.encode('utf-8', errors='ignore').decode('utf-8')
    except:
        # Fallback to ASCII if UTF-8 fails
        content = content.encode('ascii', errors='ignore').decode('ascii')
    
    # Final cleanup - remove any empty lines at start/end
    content = content.strip()
    
    return content

#Open AI api call for template content
def find_openai_placeholders(doc):
    """Find which OpenAI placeholders exist in the document"""
    try:
        # Get document content
        doc_content = doc.Content.Text
        
        # Define OpenAI placeholders to look for
        openai_placeholders = ['KEYPLAYERS', 'SEGMENTS', 'APPLICATIONS', 'TAXONOMY']
        
        # Find which ones exist in the document
        found_placeholders = []
        for placeholder in openai_placeholders:
            if placeholder in doc_content:
                found_placeholders.append(placeholder)
        
        print(f"DEBUG - Found OpenAI placeholders: {found_placeholders}")
        return found_placeholders
        
    except Exception as e:
        print(f"Error finding OpenAI placeholders: {e}")
        return []

def generate_openai_content(keyword, category, placeholders):
    """Generate content for found placeholders using OpenAI"""
    try:
        import openai
        
        # Remove "Market" from keyword for better context
        clean_keyword = keyword.replace("Market", "").strip()
        
        # Create dynamic prompt based on found placeholders
        prompt = f"""
Generate professional market research content for {clean_keyword} in the {category} domain.

Please provide content for these specific sections:
"""
        
        # Add sections based on found placeholders
        if 'KEYPLAYERS' in placeholders:
            prompt += "\n1. KEYPLAYERS: List 10-12 key market players/companies (comma-separated)"
        if 'SEGMENTS' in placeholders:
            prompt += "\n2. SEGMENTS: List market segments by type/product (comma-separated)" 
        if 'APPLICATIONS' in placeholders:
            prompt += "\n3. APPLICATIONS: List market applications/use cases (comma-separated)"
        if 'TAXONOMY' in placeholders:
            prompt += "\n4. TAXONOMY: Provide market segmentation structure (by type, application, region)"
        
        prompt += f"""

Market Context: {clean_keyword}
Industry Domain: {category}

Format your response exactly like this:
KEYPLAYERS: [content here]
SEGMENTS: [content here]  
APPLICATIONS: [content here]
TAXONOMY: [content here]

Only include sections that were requested above. Make content professional and industry-appropriate.
"""
        
        print(f"DEBUG - Making OpenAI call for {len(placeholders)} placeholders")
        
        # Make OpenAI API call
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research content generator."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        ai_content = response.choices[0].message.content
        print(f"DEBUG - OpenAI response received: {len(ai_content)} characters")
        
        return ai_content
        
    except Exception as e:
        print(f"ERROR - OpenAI API call failed: {e}")
        return None

def parse_openai_response(ai_content, placeholders):
    """Parse OpenAI response into individual placeholder content"""
    try:
        replacements = {}
        
        if not ai_content:
            return replacements
        
        # Split response into lines
        lines = ai_content.split('\n')
        
        # Parse each line for placeholder content
        for line in lines:
            if ':' in line:
                parts = line.split(':', 1)
                if len(parts) == 2:
                    placeholder = parts[0].strip()
                    content = parts[1].strip()
                    
                    # Only add if it's one of our requested placeholders
                    if placeholder in placeholders and content:
                        replacements[placeholder] = content
        
        print(f"DEBUG - Parsed {len(replacements)} placeholder replacements")
        for key, value in replacements.items():
            print(f"DEBUG - {key}: {value[:100]}...")
        
        return replacements
        
    except Exception as e:
        print(f"ERROR - Failed to parse OpenAI response: {e}")
        return {}

def replace_openai_placeholders(doc, keyword, category):
    """Main function to handle OpenAI placeholder replacement"""
    try:
        print(f"DEBUG - Starting OpenAI placeholder replacement for: {keyword}")
        
        # Step 1: Find which OpenAI placeholders exist in document
        found_placeholders = find_openai_placeholders(doc)
        
        if not found_placeholders:
            print("DEBUG - No OpenAI placeholders found in document")
            return
        
        # Step 2: Generate content using OpenAI
        ai_content = generate_openai_content(keyword, category, found_placeholders)
        
        if not ai_content:
            print("ERROR - Failed to generate OpenAI content")
            return
        
        # Step 3: Parse OpenAI response
        replacements = parse_openai_response(ai_content, found_placeholders)
        
        if not replacements:
            print("ERROR - Failed to parse OpenAI response into usable content")
            return
        
        # Step 4: Replace each placeholder in document
        for placeholder, content in replacements.items():
            print(f"DEBUG - Replacing {placeholder} with OpenAI content...")
            replace_text_preserve_formatting(doc, placeholder, content)
        
        print(f"DEBUG - OpenAI placeholder replacement completed")
        
    except Exception as e:
        print(f"ERROR - OpenAI placeholder replacement failed: {e}")


#Template code for WMR
@app.route('/custom_wmr_template_processor', methods=['GET', 'POST'])
def custom_wmr_template_processor():
    if request.method == 'POST':
        try:
            # Get the uploaded file
            template_file = request.files.get('template_file')
            
            if not template_file or template_file.filename == '':
                flash('❌ Please upload an Excel file!')
                return redirect(request.url)

            if not allowed_file(template_file.filename):
                flash('❌ Only Excel files (.xlsx, .xls) are allowed!')
                return redirect(request.url)

            # Save uploaded file
            filename = secure_filename(template_file.filename)
            input_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            template_file.save(input_path)

            # Process the WMR templates
            result = process_wmr_templates_with_excel(input_path)
            
            if result['success']:
                flash(f'✅ Successfully generated {result["files_created"]} WMR template files!')
                flash(f'📁 Files saved in WMR/Generated folder')
            else:
                flash(f'❌ Error: {result["error"]}')

            # Clean up uploaded file
            try:
                os.remove(input_path)
            except:
                pass

            return render_template('custom_wmr_templates.html')  # You'll need to create this template

        except Exception as e:
            flash(f'❌ Error processing WMR templates: {str(e)}')
            return redirect(request.url)
    
    return render_template('custom_wmr_templates.html')


def process_wmr_templates_with_excel(excel_path):
    """Process Excel file and generate WMR templates for each row"""
    try:
        print(f"\n=== PROCESSING WMR TEMPLATES ===")
        print(f"Excel file: {excel_path}")
        
        # Read the Excel file
        df = pd.read_excel(excel_path, engine='openpyxl')
        
        print(f"Excel columns: {list(df.columns)}")
        print(f"Found {len(df)} rows to process")
        
        # Verify required columns exist (no Category needed for WMR)
        required_columns = ['KEYWORD', 'PROMOBUY', 'SAMPLECOPY']
        missing_columns = [col for col in required_columns if col not in df.columns]
        
        if missing_columns:
            return {
                'success': False,
                'error': f'Missing required columns: {missing_columns}'
            }
        
        files_created = 0
        base_rpa_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        
        # Process each row
        for index, row in df.iterrows():
            try:
                keyword = str(row['KEYWORD']).strip()
                
                print(f"\n[{index+1}/{len(df)}] Processing WMR: {keyword}")
                
                # Get random WMR template
                template_path = get_random_wmr_template(base_rpa_path)
                if not template_path:
                    print(f"❌ No WMR templates found")
                    continue
                
                print(f"Selected WMR template: {os.path.basename(template_path)}")
                
                # Process the template
                success, output_path = process_single_wmr_template(template_path, row, base_rpa_path)
                
                if success:
                    print(f"✅ WMR Template generated: {os.path.basename(output_path)}")
                    files_created += 1
                else:
                    print(f"❌ Failed to generate WMR template for: {keyword}")
                
            except Exception as e:
                print(f"❌ Error processing WMR row {index+1}: {e}")
                continue
        
        return {
            'success': True,
            'files_created': files_created,
            'total_rows': len(df)
        }
        
    except Exception as e:
        print(f"Error in WMR template processing: {e}")
        return {
            'success': False,
            'error': str(e)
        }

def get_random_wmr_template(base_path):
    """Get a random WMR template from the WMR folder"""
    try:
        wmr_path = os.path.join(base_path, "WMR")
        
        if not os.path.exists(wmr_path):
            print(f"❌ WMR folder not found: {wmr_path}")
            return None
        
        # Find all .doc files in the WMR folder, excluding lock files
        doc_files = [f for f in os.listdir(wmr_path) 
                    if (f.lower().endswith('.doc') or f.lower().endswith('.docx'))
                    and not f.startswith('~$')]
        
        if not doc_files:
            print(f"❌ No .doc files found in {wmr_path}")
            return None
        
        # Select random template
        import random
        selected_template = random.choice(doc_files)
        template_path = os.path.join(wmr_path, selected_template)
        
        print(f"Random WMR template selected: {selected_template}")
        return template_path
        
    except Exception as e:
        print(f"Error getting random WMR template: {e}")
        return None

def process_single_wmr_template(template_path, row_data, base_rpa_path):
    """Process a single WMR template with row data"""
    try:
        today = datetime.today()
        
        # Generate output filename using KEYWORD
        keyword = str(row_data['KEYWORD']).strip()
        # Remove "Market" from the KEYWORD value
        keyword = keyword.replace("Market", "").strip()
        safe_keyword = "".join(c for c in keyword if c.isalnum() or c in (' ', '_')).strip()
        safe_keyword = re.sub(r'\s+', '_', safe_keyword)
        
        filename = f"WMR_{safe_keyword}_{today.year}_{today.month:02d}_{today.day:02d}.doc"
        
        # Save generated files in WMR/Generated subfolder
        generated_folder = os.path.join(base_rpa_path, "WMR", "Generated")
        if not os.path.exists(generated_folder):
            os.makedirs(generated_folder)
        
        output_path = os.path.join(generated_folder, filename)
        
        # Process the WMR template
        success = process_wmr_template_with_formatting(template_path, row_data, output_path, keyword)
        
        if success:
            return True, output_path
        else:
            return False, None
            
    except Exception as e:
        print(f"Error processing single WMR template: {e}")
        return False, None

def process_wmr_template_with_formatting(template_path, row_data, output_path, keyword):
    """Process WMR template with formatting preservation"""
    try:
        # Open Word application
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        
        # Open the template document
        doc = word.Documents.Open(template_path)
        
        # Prepare replacement data
        today = datetime.today()
        replacements = get_wmr_replacement_data(row_data, today)
        
        print(f"Applying {len(replacements)} basic WMR replacements to template")
        
        # STEP 1: Replace basic placeholders while preserving formatting
        for placeholder, replacement in replacements.items():
            if replacement:  # Only replace if we have a value
                # Check if it's 'KEYWORD' and remove "Market"
                if placeholder == 'KEYWORD':
                    replacement = replacement.replace("Market", "").strip()
                
                replace_text_preserve_formatting(doc, placeholder, str(replacement))
        
        # STEP 2: Replace OpenAI placeholders for WMR
        replace_wmr_openai_placeholders(doc, keyword)
        
        # Save the document with preserved formatting
        doc.SaveAs2(output_path, FileFormat=0)  # 0 = Word 97-2003 format
        doc.Close()
        word.Quit()
        
        return True
        
    except Exception as e:
        print(f"Error processing WMR template with formatting: {e}")
        try:
            word.Quit()
        except:
            pass
        return False

def get_wmr_replacement_data(row_data, today):
    """Get replacement data from Excel for WMR - 3 placeholders only"""
    replacements = {
        'KEYWORD': str(row_data.get('KEYWORD', '')),
        'PROMOBUY': str(row_data.get('PROMOBUY', '')), 
        'SAMPLECOPY': str(row_data.get('SAMPLECOPY', ''))
    }
    
    print(f"DEBUG - WMR Excel row keys: {list(row_data.keys())}")
    print(f"DEBUG - WMR Final replacements: {replacements}")
    
    return replacements

def replace_wmr_openai_placeholders(doc, keyword):
    """Handle OpenAI placeholder replacement for WMR"""
    try:
        print(f"DEBUG - Starting WMR OpenAI placeholder replacement for: {keyword}")
        
        # Step 1: Find which OpenAI placeholders exist in document
        found_placeholders = find_openai_placeholders(doc)  # Reuse existing function
        
        if not found_placeholders:
            print("DEBUG - No OpenAI placeholders found in WMR document")
            return
        
        # Step 2: Generate content using OpenAI (WMR context)
        ai_content = generate_wmr_openai_content(keyword, found_placeholders)
        
        if not ai_content:
            print("ERROR - Failed to generate WMR OpenAI content")
            return
        
        # Step 3: Parse OpenAI response
        replacements = parse_openai_response(ai_content, found_placeholders)  # Reuse existing function
        
        if not replacements:
            print("ERROR - Failed to parse WMR OpenAI response")
            return
        
        # Step 4: Replace each placeholder in document
        for placeholder, content in replacements.items():
            print(f"DEBUG - Replacing WMR {placeholder} with OpenAI content...")
            replace_text_preserve_formatting(doc, placeholder, content)
        
        print(f"DEBUG - WMR OpenAI placeholder replacement completed")
        
    except Exception as e:
        print(f"ERROR - WMR OpenAI placeholder replacement failed: {e}")

def generate_wmr_openai_content(keyword, placeholders):
    """Generate content for found placeholders using OpenAI for WMR"""
    try:
        import openai
        
        # Remove "Market" from keyword for better context
        clean_keyword = keyword.replace("Market", "").strip()
        
        # Create dynamic prompt based on found placeholders (WMR domain)
        prompt = f"""
Generate professional market research content for {clean_keyword} for World Market Reports (WMR).

Please provide content for these specific sections:
"""
        
        # Add sections based on found placeholders
        if 'KEYPLAYERS' in placeholders:
            prompt += "\n1. KEYPLAYERS: List 8-12 key market players/companies (comma-separated)"
        if 'SEGMENTS' in placeholders:
            prompt += "\n2. SEGMENTS: List market segments by type/product (comma-separated)" 
        if 'APPLICATIONS' in placeholders:
            prompt += "\n3. APPLICATIONS: List market applications/use cases (comma-separated)"
        if 'TAXONOMY' in placeholders:
            prompt += "\n4. TAXONOMY: Provide market segmentation structure (by type, application, region)"
        
        prompt += f"""

Market Context: {clean_keyword}
Report Type: World Market Reports (WMR)

Format your response exactly like this:
KEYPLAYERS: [content here]
SEGMENTS: [content here]  
APPLICATIONS: [content here]
TAXONOMY: [content here]

Only include sections that were requested above. Make content professional and industry-appropriate.
"""
        
        print(f"DEBUG - Making WMR OpenAI call for {len(placeholders)} placeholders")
        
        # Make OpenAI API call
        client = openai.OpenAI(api_key=OPENAI_API_KEY)
        response = client.chat.completions.create(
            model="gpt-4o-mini",
            messages=[
                {"role": "system", "content": "You are a professional market research content generator for World Market Reports."},
                {"role": "user", "content": prompt}
            ],
            temperature=0.7,
            max_tokens=800
        )
        
        ai_content = response.choices[0].message.content
        print(f"DEBUG - WMR OpenAI response received: {len(ai_content)} characters")
        
        return ai_content
        
    except Exception as e:
        print(f"ERROR - WMR OpenAI API call failed: {e}")
        return None





# Step 4 Custom OpenPR Processing
@app.route('/custom_document_processing', methods=['GET', 'POST'])
def custom_document_processing():
    if request.method == 'POST':
        try:
            flash('Custom document processing started!')
            return redirect(url_for('custom_document_processing'))
        except Exception as e:
            flash(f'Error: {str(e)}')
            return redirect(url_for('custom_document_processing'))
    
    return render_template('custom_document_processing.html', 
                          session_data={
                              'username': session.get('username', ''),
                              'email': session.get('email', ''),
                              'mobile': session.get('mobile', ''),
                              'open_pr_id': session.get('open_pr_id', ''),
                              'image_path': session.get('image_path', '')
                          })
    

@app.route('/start_custom_wmr_publishing', methods=['POST'])
def start_custom_wmr_publishing():
    """Start custom WMR publishing process"""
    try:
        # Run the selenium automation in a background thread
        def run_wmr_publishing():
            result = selenium_publishing_custom_wmr()
            print(f"WMR Publishing completed: {result}")
        
        # Start in background thread so user gets immediate response
        threading.Thread(target=run_wmr_publishing).start()
        
        return jsonify({
            'status': 'success',
            'message': 'WMR publishing started! Check console for progress updates.'
        })
        
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error starting WMR publishing: {str(e)}'
        }), 500

def selenium_publishing_custom_wmr():
    """Selenium automation to publish all WMR articles from Generated folder to OpenPR"""
    try:
        import random
        import os
        from datetime import datetime
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        from webdriver_manager.chrome import ChromeDriverManager
        import time
        import re
        
        print(f"\n=== STARTING WMR CUSTOM PUBLISHING ===")
        
        # Hardcoded form details
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.""",
            """Author of this marketing PR:
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc.""",
            """Author of this marketing PR:
Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights."""
        ]
        
        TITLE_PROMPTS = [
            "Is Booming Worldwide 2025-2032",
            "Generated Opportunities, Future Scope 2025-2032",
            "Future Business Opportunities 2025-2032",
            "Growth in Future Scope 2025-2032",
            "Is Booming So Rapidly Growth by 2032",
            "Is Booming So Rapidly with CAGR of 6.9%",
            "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
            "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
            "Set to Witness Significant Growth by 2025-2032",
            "to Witness Massive Growth by 2032",
            "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
            "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
            "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
            "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
            "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
            "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
        ]
        
        # Hardcoded user details
        author_name = "Vishwas Tiwari"
        author_email = "vishwas@coherentmarketinsights.com"
        company_name = "Coherent Market Insights"
        phone_number = "1234567890"
        article_code = "D5A-2025-QDFH8C"  # Default article code
        image_path = r"C:\Users\akshat\Desktop\code\image.jpg"  # Default image path
        website_category = "Business,Economy,Finance,Banking & Insurance"

        
        # Get all WMR articles from Generated folder
        base_rpa_custom_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        generated_folder = os.path.join(base_rpa_custom_path, "WMR", "Generated")
        
        if not os.path.exists(generated_folder):
            print(f"❌ Generated folder not found: {generated_folder}")
            return {'success': False, 'error': 'Generated folder not found'}
        
        # Get all .doc and .docx files
        article_files = []
        for file in os.listdir(generated_folder):
            if file.lower().endswith(('.doc', '.docx')) and not file.startswith('~'):
                file_path = os.path.join(generated_folder, file)
                if os.path.isfile(file_path):
                    # Get file modification time for sorting
                    mod_time = os.path.getmtime(file_path)
                    article_files.append((file_path, mod_time, file))
        
        if not article_files:
            print(f"❌ No WMR articles found in: {generated_folder}")
            return {'success': False, 'error': 'No articles found in Generated folder'}
        
        # Sort by modification time (latest first)
        article_files.sort(key=lambda x: x[1], reverse=True)
        
        print(f"✅ Found {len(article_files)} WMR articles to publish")
        for i, (path, mod_time, filename) in enumerate(article_files, 1):
            print(f"  {i}. {filename} ({datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M')})")
        
        successful_publications = 0
        failed_publications = 0
        
        # Setup Chrome driver
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        # Process each article
        for i, (article_path, mod_time, filename) in enumerate(article_files, 1):
            try:
                print(f"\n{'='*60}")
                print(f"Processing Article {i}/{len(article_files)}: {filename}")
                print(f"{'='*60}")
                
                # Extract market name from filename (fallback method)
                market_name_from_file = extract_market_name_from_filename(filename)
                
                # Read article content using text_of_press_release
                print(f"📖 Reading article content from: {filename}")
                article_content = text_of_press_release(article_path)
                
                if not article_content or len(article_content.strip()) < 100:
                    print(f"❌ Article content too short or empty, skipping: {filename}")
                    failed_publications += 1
                    continue
                
                # Generate article title
                random_prompt = random.choice(TITLE_PROMPTS)
                article_title = f"{market_name_from_file} {random_prompt}"
                print(f"📝 Generated title: {article_title}")
                
                # Start Selenium automation
                print(f"🚀 Starting Selenium automation for: {market_name_from_file}")
                cService = Service(executable_path=chromedriver_path)
                driver = webdriver.Chrome(service=cService, options=options)
                driver.get('https://www.openpr.com/')
                
                # Handle cookie consent
                try:
                    reject = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
                    )
                    reject.click()
                    print("✅ Cookie consent handled")
                except:
                    print("⚠️ Cookie consent button not found or already handled")
                
                # Navigate to submit page
                submit = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
                )
                submit.click()
                print("✅ Navigated to submit page")
                
                # Enter article code
                input_box = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
                )
                input_box.clear()
                input_box.send_keys(article_code)
                print(f"✅ Entered article code: {article_code}")
                
                # Submit code
                try:
                    submit2 = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
                    )
                    submit2.click()
                except:
                    submit2 = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
                    )
                    submit2.click()
                print("✅ Article code submitted")
                
                # Fill form fields
                name = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
                )
                name.send_keys(author_name)
                
                email = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
                )
                email.clear()
                email.send_keys(author_email)
                
                number = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
                )
                number.clear()
                number.send_keys(phone_number)
                
                ComName = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
                )
                ComName.clear()
                ComName.send_keys(company_name)
                
                s1 = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
                )
                s1.click()
                print("✅ Basic form fields filled")
                
                # Handle category selection
                Category_element = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                )
                time.sleep(1)
                
                select_obj = Select(Category_element)
                try:
                    select_obj.select_by_visible_text(website_category)
                    print(f"✅ Selected category: {website_category}")
                except Exception as e:
                    print(f"⚠️ Category selection failed, using fallback: {e}")
                    select_obj.select_by_index(1)  # Select first real option
                
                # Fill title
                title = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
                )
                title.clear()
                title.send_keys(article_title)
                print(f"✅ Entered title: {len(article_title)} characters")
                
                # Fill article content (main text)
                text = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
                )
                text.clear()
                text.send_keys(article_content)
                print(f"✅ Entered article content: {len(article_content)} characters")
                
                # Fill about section
                about = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
                )
                about.clear()
                contact_info = """Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837"""
                about.send_keys(contact_info)
                
                # Fill address section with random author
                address = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
                )
                address.clear()
                random_author = random.choice(AUTHOR_DESCRIPTIONS)
                address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."
                address.send_keys(address_content)
                print("✅ About and contact information filled")
                
                # Upload image
                if os.path.exists(image_path):
                    image = WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
                    )
                    image.clear()
                    image.send_keys(image_path)
                    print(f"✅ Image uploaded: {image_path}")
                else:
                    print(f"⚠️ Image file not found: {image_path}")
                
                # Fill caption
                caption = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
                )
                caption.clear()
                caption.send_keys(f"Market analysis image for {market_name_from_file}")
                
                # Fill notes
                notes = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
                )
                notes.clear()
                notes.send_keys(f"Comprehensive market research report on {market_name_from_file} with detailed analysis and forecasts.")
                
                # Agree to terms
                tick1 = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
                )
                tick1.click()
                
                tick2 = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
                )
                tick2.click()
                print("✅ Terms and conditions accepted")
                
                # Submit form
                final = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
                )
                final.click()
                print("🚀 Form submitted!")
                
                # Wait for submission to complete
                time.sleep(5)
                
                print(f"✅ Successfully published: {filename}")
                successful_publications += 1
                
                # Close browser for this article
                # driver.quit()
                
                # Add delay between articles to avoid rate limiting
                if i < len(article_files):  # Don't wait after the last article
                    print(f"⏳ Waiting 30 seconds before next article...")
                    time.sleep(30)
                
            except Exception as e:
                print(f"❌ Error publishing {filename}: {e}")
                failed_publications += 1
                try:
                    driver.quit()
                except:
                    pass
                continue
        
        # Final summary
        print(f"\n{'='*60}")
        print("WMR PUBLISHING SUMMARY")
        print(f"{'='*60}")
        print(f"✅ Successfully published: {successful_publications}")
        print(f"❌ Failed publications: {failed_publications}")
        print(f"📊 Total articles processed: {len(article_files)}")
        print(f"{'='*60}")
        
        return {
            'success': True,
            'total_articles': len(article_files),
            'successful_publications': successful_publications,
            'failed_publications': failed_publications
        }
        
    except Exception as e:
        print(f"❌ Error in WMR publishing automation: {e}")
        return {'success': False, 'error': str(e)}


def extract_market_name_from_filename(filename):
    """Extract market name from WMR filename"""
    try:
        # Remove file extension
        name_without_ext = os.path.splitext(filename)[0]
        
        # Remove WMR prefix and date suffix
        # Pattern: WMR_Market_Name_2025_08_01
        parts = name_without_ext.split('_')
        
        if len(parts) >= 4 and parts[0] == 'WMR':
            # Remove WMR prefix and last 3 parts (year, month, day)
            market_parts = parts[1:-3]
            market_name = ' '.join(market_parts)
            
            # Add "Market" back if not present
            if not market_name.lower().endswith('market'):
                market_name += ' Market'
                
            return market_name
        else:
            # Fallback: use filename as is
            return name_without_ext.replace('_', ' ')
            
    except Exception as e:
        print(f"Error extracting market name from {filename}: {e}")
        return filename.replace('_', ' ')


def text_of_press_release(doc_path):
    """Extract and format press release text from Word document"""
    import win32com.client
    import re
    
    # Open Word application
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False  # Run in the background

    # Open the .doc file (adjust the file path if necessary)
    doc2 = word.Documents.Open(doc_path)

    # Extract the entire text from the document
    doc_text = doc2.Content.Text

    # Remove the first line from the document text
    lines = doc_text.splitlines()
    if len(lines) > 1:
        doc_text = '\n'.join(lines[1:])

    # Define the headings for which you want to add line breaks
    headings = [
        "➤Market Size and Overview",
        "➤Actionable Insights",
        "➤Actionable insights",
        "➤Growth factors",
        "➤Growth Factors",
        "➤Market trends",
        "➤Market Trends",
        "➤Key takeaways ",
        "➤Key Takeaways",
        "➤Market Segment and Regional Coverage ",
        "➤Market segment and regional coverage",
        "➤Key players",
        "➤Key Players",
        "➤Competitive Strategies and Outcomes",
        "❓ Frequently Asked Questions",
        "❓ Frequently asked questions"
    ]

    # FIXED: Add a line space AFTER each heading (not before and after)
    for heading in headings:
        doc_text = doc_text.replace(heading, f"{heading}\n")

    # Define the regex pattern for URLs
    url_pattern = re.compile(r"(https?://[^\s]+)")
    
    # Define regex patterns for FAQ questions (numbered questions and roman numerals)
    faq_pattern_numbers = re.compile(r"^\d+\.\s")  # Matches "1. ", "2. ", etc.
    faq_pattern_roman = re.compile(r"^[ivxlcdmIVXLCDM]+\.\s")  # Matches "i. ", "ii. ", "I. ", "II. ", etc.
    
    # Define regex pattern for CTA links (➔)
    cta_pattern = re.compile(r"^➔")  # Matches lines starting with ➔

    # Split the text into lines
    lines = doc_text.splitlines()
    processed_lines = []

    # Iterate over each line
    for i, line in enumerate(lines):
        line_stripped = line.strip()
        
        # Skip empty lines in processing, we'll add them strategically
        if not line_stripped:
            continue
            
        # Check if this is a CTA line
        is_cta = cta_pattern.match(line_stripped)
        
        # Check if previous line was a CTA (for adjacent CTA handling)
        prev_was_cta = False
        if processed_lines:
            last_non_empty = None
            for prev_line in reversed(processed_lines):
                if prev_line.strip():
                    last_non_empty = prev_line.strip()
                    break
            if last_non_empty and cta_pattern.match(last_non_empty):
                prev_was_cta = True
        
        # Check if this line is a heading (starts with ➤ or ❓)
        is_heading = line_stripped.startswith('➤') or line_stripped.startswith('❓')
        
        # If a line contains a URL, add space before and after the URL
        if url_pattern.search(line):
            # Add space before (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            processed_lines.append('')  # Add space after
            
        # If a line is an FAQ question (starts with number or roman numeral), add space before it
        elif faq_pattern_numbers.match(line_stripped) or faq_pattern_roman.match(line_stripped):
            # Add space before FAQ question (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this is a CTA line
        elif is_cta:
            # Add space before CTA (unless previous was also CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            
        # If this line is a heading (starts with ➤ or ❓)
        elif is_heading:
            # Add space before heading (unless previous was CTA)
            if not prev_was_cta and processed_lines and processed_lines[-1].strip():
                processed_lines.append('')
            processed_lines.append(line)
            # FIXED: Add space AFTER heading
            processed_lines.append('')
            
        else:
            # Regular content line
            processed_lines.append(line)

    # Join the processed lines back into a single string
    chunk = "\n".join(processed_lines)
    
    # Clean up multiple consecutive empty lines (replace with single empty line)
    chunk = re.sub(r'\n\s*\n\s*\n+', '\n\n', chunk)

    # Close the document
    doc2.Close()
    word.Quit()

    # Return the processed content
    return chunk

@app.route('/start_custom_ai_publishing', methods=['POST'])
def start_custom_ai_publishing():
    try:
        # Add your AI publishing logic here
        # This could involve reading AI-generated content files
        # and publishing them to OpenPR
        
        return jsonify({
            'status': 'success',
            'message': 'AI content publishing started successfully!'
        })
    except Exception as e:
        return jsonify({
            'status': 'error',
            'message': f'Error starting AI publishing: {str(e)}'
        })

def selenium_publishing_custom_cmi():
    """Selenium automation to publish all WMR articles from Generated folder to OpenPR"""
    try:
        import random
        import os
        from datetime import datetime
        from selenium import webdriver
        from selenium.webdriver.chrome.service import Service
        from selenium.webdriver.common.by import By
        from selenium.webdriver.chrome.options import Options
        from selenium.webdriver.support.ui import WebDriverWait
        from selenium.webdriver.support import expected_conditions as EC
        from selenium.webdriver.support.ui import Select
        from webdriver_manager.chrome import ChromeDriverManager
        import time
        import re
        
        print(f"\n=== STARTING WMR CUSTOM PUBLISHING ===")
        
        # Hardcoded form details
        AUTHOR_DESCRIPTIONS = [
            """Author of this marketing PR:
Ravina Pandya, Content Writer, has a strong foothold in the market research industry. She specializes in writing well-researched articles from different industries, including food and beverages, information and technology, healthcare, chemical and materials, etc.""",
            """Author of this marketing PR:
Money Singh is a seasoned content writer with over four years of experience in the market research sector. Her expertise spans various industries, including food and beverages, biotechnology, chemical and materials, defense and aerospace, consumer goods, etc.""",
            """Author of this marketing PR:
Alice Mutum is a seasoned senior content editor at Coherent Market Insights, leveraging extensive expertise gained from her previous role as a content writer. With seven years in content development, Alice masterfully employs SEO best practices and cutting-edge digital marketing strategies to craft high-ranking, impactful content. As an editor, she meticulously ensures flawless grammar and punctuation, precise data accuracy, and perfect alignment with audience needs in every research report. Alice's dedication to excellence and her strategic approach to content make her an invaluable asset in the world of market insights."""
        ]
        
        TITLE_PROMPTS = [
            "Is Booming Worldwide 2025-2032",
            "Generated Opportunities, Future Scope 2025-2032",
            "Future Business Opportunities 2025-2032",
            "Growth in Future Scope 2025-2032",
            "Is Booming So Rapidly Growth by 2032",
            "Is Booming So Rapidly with CAGR of 6.9%",
            "An Analysis of Size, Shares, Business Growth, and Upcoming Trends Forecast 2025-2032",
            "2025-2032 Emerging Trends in Industry Dynamics, Size Insights, Share, and Future Growth",
            "Set to Witness Significant Growth by 2025-2032",
            "to Witness Massive Growth by 2032",
            "Size, Share 2025 Analysis of Rising Business Opportunities with Prominent Investment, Forecast to 2032",
            "by Trends, Dynamic Innovation in Technology and 2032 Forecast, Opportunities, and Challenges, Trends",
            "Size 2025 Emerging Demands, Share, Trends, Futuristic Opportunity, Share and Forecast To 2032",
            "Size, Share Growth Status, Emerging Technology, Key Players, Industry Challenges, and Forecast till 2032",
            "Insights 2025-2032: Global Expansion, Revenue Trends, and Strategic Growth Plans",
            "Outlook 2025-2032: Growth Drivers, Share, And Trends As Revealed In New Report"
        ]
        
        # Hardcoded user details
        author_name = "Vishwas Tiwari"
        author_email = "vishwas@coherentmarketinsights.com"
        company_name = "Coherent Market Insights"
        phone_number = "1234567890"
        article_code = "D5A-2025-QDFH8C"  # Default article code
        image_path = r"C:\Users\akshat\Desktop\code\image.jpg"  # Default image path
        website_category = "Business,Economy,Finance,Banking & Insurance"

        
        # Get all WMR articles from Generated folder
        base_rpa_custom_path = os.path.join(os.path.expanduser("~"), "Desktop", "RPA")
        generated_folder = os.path.join(base_rpa_custom_path, "WMR", "Generated")
        
        if not os.path.exists(generated_folder):
            print(f"❌ Generated folder not found: {generated_folder}")
            return {'success': False, 'error': 'Generated folder not found'}
        
        # Get all .doc and .docx files
        article_files = []
        for file in os.listdir(generated_folder):
            if file.lower().endswith(('.doc', '.docx')) and not file.startswith('~'):
                file_path = os.path.join(generated_folder, file)
                if os.path.isfile(file_path):
                    # Get file modification time for sorting
                    mod_time = os.path.getmtime(file_path)
                    article_files.append((file_path, mod_time, file))
        
        if not article_files:
            print(f"❌ No WMR articles found in: {generated_folder}")
            return {'success': False, 'error': 'No articles found in Generated folder'}
        
        # Sort by modification time (latest first)
        article_files.sort(key=lambda x: x[1], reverse=True)
        
        print(f"✅ Found {len(article_files)} WMR articles to publish")
        for i, (path, mod_time, filename) in enumerate(article_files, 1):
            print(f"  {i}. {filename} ({datetime.fromtimestamp(mod_time).strftime('%Y-%m-%d %H:%M')})")
        
        successful_publications = 0
        failed_publications = 0
        
        # Setup Chrome driver
        chromedriver_path = ChromeDriverManager().install()
        options = Options()
        options.add_argument("--start-maximized")
        
        # Process each article
        for i, (article_path, mod_time, filename) in enumerate(article_files, 1):
            try:
                print(f"\n{'='*60}")
                print(f"Processing Article {i}/{len(article_files)}: {filename}")
                print(f"{'='*60}")
                
                # Extract market name from filename (fallback method)
                market_name_from_file = extract_market_name_from_filename(filename)
                
                # Read article content using text_of_press_release
                print(f"📖 Reading article content from: {filename}")
                article_content = text_of_press_release(article_path)
                
                if not article_content or len(article_content.strip()) < 100:
                    print(f"❌ Article content too short or empty, skipping: {filename}")
                    failed_publications += 1
                    continue
                
                # Generate article title
                random_prompt = random.choice(TITLE_PROMPTS)
                article_title = f"{market_name_from_file} {random_prompt}"
                print(f"📝 Generated title: {article_title}")
                
                # Start Selenium automation
                print(f"🚀 Starting Selenium automation for: {market_name_from_file}")
                cService = Service(executable_path=chromedriver_path)
                driver = webdriver.Chrome(service=cService, options=options)
                driver.get('https://www.openpr.com/')
                
                # Handle cookie consent
                try:
                    reject = WebDriverWait(driver, 10).until(
                        EC.element_to_be_clickable((By.XPATH, '//*[@id="cmpbntnotxt"]'))
                    )
                    reject.click()
                    print("✅ Cookie consent handled")
                except:
                    print("⚠️ Cookie consent button not found or already handled")
                
                # Navigate to submit page
                submit = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="navbarText"]/ul/li[3]/a'))
                )
                submit.click()
                print("✅ Navigated to submit page")
                
                # Enter article code
                input_box = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="code"]'))
                )
                input_box.clear()
                input_box.send_keys(article_code)
                print(f"✅ Entered article code: {article_code}")
                
                # Submit code
                try:
                    submit2 = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(5) > div > form > button'))
                    )
                    submit2.click()
                except:
                    submit2 = WebDriverWait(driver, 5).until(
                        EC.element_to_be_clickable((By.CSS_SELECTOR, '#main > div > div > div:nth-child(6) > div > form > button'))
                    )
                    submit2.click()
                print("✅ Article code submitted")
                
                # Fill form fields
                name = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[1]/div/input'))
                )
                name.send_keys(author_name)
                
                email = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[2]/div/input'))
                )
                email.clear()
                email.send_keys(author_email)
                
                number = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[4]/div/input'))
                )
                number.clear()
                number.send_keys(phone_number)
                
                ComName = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="archivnmfield"]'))
                )
                ComName.clear()
                ComName.send_keys(company_name)
                
                s1 = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="popup-archiv"]/div/a[1]'))
                )
                s1.click()
                print("✅ Basic form fields filled")
                
                # Handle category selection
                Category_element = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/div[6]/div/select'))
                )
                time.sleep(1)
                
                select_obj = Select(Category_element)
                try:
                    select_obj.select_by_visible_text(website_category)
                    print(f"✅ Selected category: {website_category}")
                except Exception as e:
                    print(f"⚠️ Category selection failed, using fallback: {e}")
                    select_obj.select_by_index(1)  # Select first real option
                
                # Fill title
                title = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[7]/div/input'))
                )
                title.clear()
                title.send_keys(article_title)
                print(f"✅ Entered title: {len(article_title)} characters")
                
                # Fill article content (main text)
                text = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="inhalt"]'))
                )
                text.clear()
                text.send_keys(article_content)
                print(f"✅ Entered article content: {len(article_content)} characters")
                
                # Fill about section
                about = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[9]/div/textarea'))
                )
                about.clear()
                contact_info = """Contact Us:

Mr. Shah
Coherent Market Insights
533 Airport Boulevard,
Suite 400, Burlingame,
CA 94010, United States
US: + 12524771362
UK: +442039578553
AUS: +61-8-7924-7805
India: +91-848-285-0837"""
                about.send_keys(contact_info)
                
                # Fill address section with random author
                address = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[10]/div/textarea'))
                )
                address.clear()
                random_author = random.choice(AUTHOR_DESCRIPTIONS)
                address_content = f"{random_author}\n\nAbout Us:\nCoherent Market Insights leads into data and analytics, audience measurement, consumer behaviors, and market trend analysis. From shorter dispatch to in-depth insights, CMI has exceled in offering research, analytics, and consumer-focused shifts for nearly a decade. With cutting-edge syndicated tools and custom-made research services, we empower businesses to move in the direction of growth. We are multifunctional in our work scope and have 450+ seasoned consultants, analysts, and researchers across 26+ industries spread out in 32+ countries."
                address.send_keys(address_content)
                print("✅ About and contact information filled")
                
                # Upload image
                if os.path.exists(image_path):
                    image = WebDriverWait(driver, 10).until(
                        EC.presence_of_element_located((By.XPATH, '//*[@id="bild"]'))
                    )
                    image.clear()
                    image.send_keys(image_path)
                    print(f"✅ Image uploaded: {image_path}")
                else:
                    print(f"⚠️ Image file not found: {image_path}")
                
                # Fill caption
                caption = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[12]/div/input'))
                )
                caption.clear()
                caption.send_keys(f"Market analysis image for {market_name_from_file}")
                
                # Fill notes
                notes = WebDriverWait(driver, 10).until(
                    EC.presence_of_element_located((By.XPATH, '//*[@id="formular"]/div[2]/div[13]/div/textarea'))
                )
                notes.clear()
                notes.send_keys(f"Comprehensive market research report on {market_name_from_file} with detailed analysis and forecasts.")
                
                # Agree to terms
                tick1 = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-agb"]'))
                )
                tick1.click()
                
                tick2 = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="input-ds"]'))
                )
                tick2.click()
                print("✅ Terms and conditions accepted")
                
                # Submit form
                final = WebDriverWait(driver, 10).until(
                    EC.element_to_be_clickable((By.XPATH, '//*[@id="formular"]/div[2]/button'))
                )
                final.click()
                print("🚀 Form submitted!")
                
                # Wait for submission to complete
                time.sleep(5)
                
                print(f"✅ Successfully published: {filename}")
                successful_publications += 1
                
                # Close browser for this article
                # driver.quit()
                
                # Add delay between articles to avoid rate limiting
                if i < len(article_files):  # Don't wait after the last article
                    print(f"⏳ Waiting 30 seconds before next article...")
                    time.sleep(30)
                
            except Exception as e:
                print(f"❌ Error publishing {filename}: {e}")
                failed_publications += 1
                try:
                    driver.quit()
                except:
                    pass
                continue
        
        # Final summary
        print(f"\n{'='*60}")
        print("WMR PUBLISHING SUMMARY")
        print(f"{'='*60}")
        print(f"✅ Successfully published: {successful_publications}")
        print(f"❌ Failed publications: {failed_publications}")
        print(f"📊 Total articles processed: {len(article_files)}")
        print(f"{'='*60}")
        
        return {
            'success': True,
            'total_articles': len(article_files),
            'successful_publications': successful_publications,
            'failed_publications': failed_publications
        }
        
    except Exception as e:
        print(f"❌ Error in WMR publishing automation: {e}")
        return {'success': False, 'error': str(e)}






if __name__ == '__main__':
    import webbrowser
    webbrowser.open('http://127.0.0.1:5000/')
    app.run(debug=True, host='0.0.0.0', port=5000)
